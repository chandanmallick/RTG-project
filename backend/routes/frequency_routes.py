# frequency_routes.py
import json
import base64
import io
import os
import uuid
import shutil
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime, timedelta, time
import requests
import pandas as pd
import numpy as np

# Non-interactive matplotlib backend
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

from docx import Document
import docx.shared

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from fastapi import APIRouter, File, UploadFile, Query, Form, Body
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
from services.db_handler import MongoService

router = APIRouter(
    prefix="/api/frequency",
    tags=["Frequency Report"]
)

SSE_REPORT_JOBS = {}

# ──────────────────────────────────────────────────────────────
# TYPE-SAFE FORMATTING & METRIC RESOLUTION HELPERS
# ──────────────────────────────────────────────────────────────

def safe_float(val):
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    try:
        val_str = str(val).strip()
        if val_str in ["—", "", "None", "null"]:
            return None
        return float(val_str.replace("%", "").strip())
    except ValueError:
        return None

def safe_round(val, decimals=0):
    f_val = safe_float(val)
    if f_val is None:
        return "—"
    return round(f_val, decimals)

def safe_format_mw(val):
    f_val = safe_float(val)
    if f_val is None:
        return "—"
    return f"{f_val:.0f} MW"

def safe_format_hz(val):
    f_val = safe_float(val)
    if f_val is None:
        return "—"
    return f"{f_val:.3f} Hz"

def safe_format_pct(val):
    f_val = safe_float(val)
    if f_val is None:
        return "—"
    return f"{f_val:.1f}%"

def get_stat(row, key):
    stats = row.get("statistics", {}) or {}
    aliases = {
        "max_od_freq": ["freq_at_max_od", "max_od_freq", "max_od_freq_str"],
        "over_drawal_pct": ["od_duration_pct", "over_drawal_pct"],
        "under_drawal_pct": ["helping_duration_pct", "under_drawal_pct"],
    }
    
    val = row.get(key)
    if val is not None:
        return val
    val = stats.get(key)
    if val is not None:
        return val
        
    if key in aliases:
        for alias in aliases[key]:
            val = row.get(alias)
            if val is not None:
                return val
            val = stats.get(alias)
            if val is not None:
                return val
                
    for canonical, alias_list in aliases.items():
        if key in alias_list:
            for alias in [canonical] + alias_list:
                if alias == key:
                    continue
                val = row.get(alias)
                if val is not None:
                    return val
                val = stats.get(alias)
                if val is not None:
                    return val
    return None

def normalize_event_type(value):
    text = str(value or "low").strip().lower()
    if text in ["high", "hf", "high_frequency", "high-frequency"]:
        return "high"
    return "low"

def event_config(event_type):
    mode = normalize_event_type(event_type)
    if mode == "high":
        return {
            "event_type": "high",
            "label": "High Frequency Operation",
            "threshold_label": "Freq>50.05",
            "threshold": 50.05,
            "is_event_freq": lambda freq: freq > 50.05,
            "state_max_key": "max_ud",
            "state_max_label": "Max UD",
            "state_max_compare": lambda dev, current: dev < current,
            "state_max_initial": 999999.0,
        }
    return {
        "event_type": "low",
        "label": "Low Frequency Operation",
        "threshold_label": "Freq<49.9",
        "threshold": 49.9,
        "is_event_freq": lambda freq: freq < 49.9,
        "state_max_key": "max_od",
        "state_max_label": "Max OD",
        "state_max_compare": lambda dev, current: dev > current,
        "state_max_initial": -999999.0,
    }

def lookup_capacity_on_bar(db, entity_list, date_strings):
    if not date_strings:
        date_strings = []
    snapshots = []
    for d in date_strings:
        doc = db.rtg_dashboard_collection.find_one(
            {"snapshot_date": d},
            sort=[("snapshot_time", -1)]
        )
        if doc:
            snapshots.append(doc)
    if not snapshots:
        latest = db.rtg_dashboard_collection.find_one({}, sort=[("snapshot_time", -1)])
        if latest:
            snapshots.append(latest)

    by_id = {}
    by_name = {}
    for doc in snapshots:
        for item in doc.get("data", []) or []:
            cap = safe_float(item.get("cap_on_bar"))
            if cap is None:
                continue
            for key in [item.get("plant_id"), item.get("rtg_plant_id")]:
                if key:
                    by_id[str(key).strip()] = cap
            name = str(item.get("plant_name") or "").strip().upper()
            if name:
                by_name[name] = cap

    result = {}
    for e in entity_list or []:
        pid = str(e.get("plant_id") or "").strip()
        rtg_pid = str(e.get("rtg_plant_id") or "").strip()
        name = str(e.get("plant_name") or e.get("STAGE_NAME") or "").strip().upper()
        cap = by_id.get(rtg_pid) or by_id.get(pid) or by_name.get(name)
        if cap is not None:
            result[pid] = cap
            if rtg_pid:
                result[rtg_pid] = cap
    return result

def compute_frequency_statistics(scada_freqs, scada_dts, entity_dev, is_state, event_type):
    cfg = event_config(event_type)
    total_count = len(scada_freqs)
    pos_count = 0
    neg_count = 0
    state_extreme_val = cfg["state_max_initial"]
    state_extreme_time = None
    state_extreme_freq = 50.0

    if entity_dev is None:
        return {
            "positive_pct": None,
            "negative_pct": None,
            "state_extreme": None,
            "state_extreme_time": "",
            "state_extreme_freq": None,
        }

    for idx, freq in enumerate(scada_freqs):
        dev_val = entity_dev[idx]
        if cfg["is_event_freq"](freq):
            if dev_val > 0:
                pos_count += 1
            elif dev_val < 0:
                neg_count += 1
        if is_state and cfg["state_max_compare"](dev_val, state_extreme_val):
            state_extreme_val = dev_val
            state_extreme_time = scada_dts[idx]
            state_extreme_freq = freq

    return {
        "positive_pct": (pos_count / total_count * 100) if total_count > 0 else 0.0,
        "negative_pct": (neg_count / total_count * 100) if total_count > 0 else 0.0,
        "state_extreme": state_extreme_val if state_extreme_time is not None else None,
        "state_extreme_time": state_extreme_time.strftime('%d-%m-%y %H:%M') if state_extreme_time is not None else "",
        "state_extreme_freq": state_extreme_freq if state_extreme_time is not None else None,
    }

# ──────────────────────────────────────────────────────────────
# SETTINGS & CONFIGURATION
# ──────────────────────────────────────────────────────────────

def get_rtg_schedule_url():
    db = MongoService()
    record = db.pipeline_config_collection.find_one({"config_type": "RTG"})
    if record and "rtg_schedule_url" in record:
        return record["rtg_schedule_url"]
    
    default_url = "https://rtgapi.grid-india.in/sendData/wbes-data/"
    db.pipeline_config_collection.update_one(
        {"config_type": "RTG"},
        {"$set": {"rtg_schedule_url": default_url}},
        upsert=True
    )
    return default_url

def get_rtg_scada_url():
    db = MongoService()
    record = db.pipeline_config_collection.find_one({"config_type": "RTG"})
    if record and "rtg_scada_url" in record:
        return record["rtg_scada_url"]
    
    default_url = "https://rtgapi.grid-india.in/sendData/scada-data/"
    db.pipeline_config_collection.update_one(
        {"config_type": "RTG"},
        {"$set": {"rtg_scada_url": default_url}},
        upsert=True
    )
    return default_url

@router.get("/settings")
async def get_settings():
    sch_url = get_rtg_schedule_url()
    scada_url = get_rtg_scada_url()
    return {
        "success": True,
        "rtg_schedule_url": sch_url,
        "rtg_scada_url": scada_url
    }

@router.put("/settings")
async def update_settings(payload: dict):
    sch_url = payload.get("rtg_schedule_url", "").strip()
    scada_url = payload.get("scada_url", "").strip()
    if not sch_url or not scada_url:
        return {"success": False, "error": "URLs cannot be empty"}
    
    db = MongoService()
    db.pipeline_config_collection.update_one(
        {"config_type": "RTG"},
        {"$set": {"rtg_schedule_url": sch_url, "rtg_scada_url": scada_url}},
        upsert=True
    )
    return {"success": True, "message": "Settings updated successfully"}

@router.get("/check-rtg-status")
async def check_rtg_status(start_time: str, end_time: str):
    try:
        st = datetime.fromisoformat(start_time.replace("Z", ""))
        et = datetime.fromisoformat(end_time.replace("Z", ""))
        
        unique_dates = []
        curr = st.date()
        while curr <= et.date():
            unique_dates.append(curr.isoformat())
            curr += timedelta(days=1)
            
        db = MongoService()
        status_by_date = {}
        all_available = True
        
        for d in unique_dates:
            doc = db.rtg_dashboard_collection.find_one(
                {"snapshot_date": d},
                sort=[("snapshot_time", -1)]
            )
            if doc and doc.get("record_count", 0) > 0:
                records = doc.get("data", [])
                has_actuals = any(float(r.get("actual_gen") or r.get("actual_gen_derived") or 0) > 0 for r in records)
                if has_actuals:
                    status_by_date[d] = True
                else:
                    status_by_date[d] = False
                    all_available = False
            else:
                status_by_date[d] = False
                all_available = False
                
        return {
            "success": True,
            "all_available": all_available,
            "status_by_date": status_by_date,
            "message": "Actual data is reporting on RTG. Upload only Frequency Excel file." if all_available else "Actual data is NOT fully reporting on RTG. Both SCADA Actuals & Frequency data upload are required."
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

# ──────────────────────────────────────────────────────────────
# GET /api/frequency/plant-mapping
# ──────────────────────────────────────────────────────────────
@router.get("/plant-mapping")
async def get_plant_mapping():
    db = MongoService()
    data = list(db.map_collection.find({}, {"_id": 0}))
    for row in data:
        for key in SCADA_KEY_FIELDS:
            row[key] = normalize_scada_key_value(row.get(key, ""))
        for key in ["plant_id", "STAGE_ID", "rtg_plant_id"]:
            if key in row:
                row[key] = clean_mapping_value(key, row.get(key, ""))
    return {"success": True, "data": data}

# ──────────────────────────────────────────────────────────────
# PUT /api/frequency/plant-mapping
# ──────────────────────────────────────────────────────────────
@router.put("/plant-mapping")
async def update_plant_mapping(payload = Body(...)):
    db = MongoService()
    rows = payload.get("rows") if isinstance(payload, dict) else payload
    if not isinstance(rows, list):
        return {"success": False, "error": "Expected mapping rows as a JSON array.", "received_type": type(payload).__name__}
    updated = 0
    matched = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        plant_id = row.get("plant_id")
        if plant_id in [None, ""]:
            continue
        stage_id = row.get("STAGE_ID")
        plant_id_str = str(plant_id).strip()
        stage_id_str = str(stage_id).strip() if stage_id not in [None, ""] else ""
        selector = {"plant_id": {"$in": [plant_id, plant_id_str]}}
        if stage_id not in [None, ""]:
            selector["STAGE_ID"] = {"$in": [stage_id, stage_id_str]}

        set_doc = {
            "wbes_name":          row.get("wbes_name", ""),
            "wbes_acronym":       row.get("wbes_acronym", ""),
            "scada_key":          normalize_scada_key_value(row.get("scada_key", "")),
            "scada_header":       row.get("scada_header", ""),
            "scada_schedule_key": normalize_scada_key_value(row.get("scada_schedule_key", "")),
            "scada_schedule_header": row.get("scada_schedule_header", ""),
            "scada_dc_key":       normalize_scada_key_value(row.get("scada_dc_key", "")),
            "scada_dc_header":    row.get("scada_dc_header", ""),
            "schedule_source":    row.get("schedule_source", "RTG"),
            "dc_source":          row.get("dc_source", "RTG"),
            "actual_source":      row.get("actual_source", "RTG"),
            "type":               row.get("type", "IPP"),
            "rtg_plant_id":       clean_mapping_value("rtg_plant_id", row.get("rtg_plant_id", "")),
            "is_state":           row.get("is_state", False),
            "is_frequency":       row.get("is_frequency", False),
            "mapping_updated_at": datetime.utcnow().isoformat(),
        }

        result = db.map_collection.update_one(
            selector,
            {"$set": set_doc},
            upsert=False
        )
        if result.matched_count == 0 and stage_id not in [None, ""]:
            result = db.map_collection.update_one(
                {"plant_id": {"$in": [plant_id, plant_id_str]}},
                {"$set": set_doc},
                upsert=False
        )
        matched += result.matched_count
        updated += result.modified_count
    return {"success": True, "updated": updated, "matched": matched, "requested": len(rows)}

# ──────────────────────────────────────────────────────────────
# DATA ACQUISITION & PROCESSING UTILS
# ──────────────────────────────────────────────────────────────

def get_legacy_session_no_verify():
    import ssl
    import urllib3
    class CustomHttpAdapterNoVerify(requests.adapters.HTTPAdapter):
        def __init__(self, ssl_context=None, **kwargs):
            self.ssl_context = ssl_context
            super().__init__(**kwargs)

        def init_poolmanager(self, connections, maxsize, block=False):
            self.poolmanager = urllib3.poolmanager.PoolManager(
                num_pools=connections,
                maxsize=maxsize,
                block=block,
                ssl_context=self.ssl_context
            )

    ctx = ssl.create_default_context(ssl.Purpose.SERVER_AUTH)
    ctx.options |= 0x4  # OP_LEGACY_SERVER_CONNECT
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    session = requests.session()
    session.mount('https://', CustomHttpAdapterNoVerify(ctx))
    return session

def log_api_hit(api_type: str, date_str: str, target: str, url: str, status: str, status_code: int, error_message: str = "", data_count: int = 0, data_saved: bool = False):
    try:
        db = MongoService()
        db.pipeline_log_collection.insert_one({
            "timestamp": datetime.utcnow().isoformat(),
            "event_type": "api_fetch",
            "api_type": api_type,
            "date": date_str,
            "target": target,
            "url": url,
            "status": status,
            "status_code": status_code,
            "error_message": error_message,
            "data_count": data_count,
            "data_saved": data_saved
        })
    except Exception as e:
        print("Error logging API hit to Mongo:", e)

RAW_DATA_COLLECTION = "frequency_event_raw_data"
EVENT_COLLECTION = "frequency_events"

def normalize_series(values, length=96):
    result = []
    for value in values or []:
        parsed = safe_float(value)
        result.append(parsed if parsed is not None else 0.0)
    if len(result) < length:
        result.extend([0.0] * (length - len(result)))
    return result[:length]

def get_unique_date_strings(start_dt: datetime, end_dt: datetime):
    dates = []
    curr = start_dt.date()
    while curr <= end_dt.date():
        dates.append(curr.isoformat())
        curr += timedelta(days=1)
    return dates

def block_index_for_time(dt_value: datetime, series_len: int):
    if series_len <= 0:
        return 0
    block_size = max(1, 1440 // series_len)
    minute_of_day = dt_value.hour * 60 + dt_value.minute
    return min(minute_of_day // block_size, series_len - 1)

def merge_event_raw_data(db, date_iso: str, plant_id: str = "", wbes_name: str = "", set_fields: Optional[dict] = None):
    dd_mm_yyyy, yyyy_mm_dd = get_date_formats(date_iso)
    query = {
        "date": yyyy_mm_dd,
        "$or": [
            {"plant_id": plant_id or ""},
            {"wbes_name": wbes_name or ""},
            {"wbes_acronym": wbes_name or ""},
        ],
    }
    if not plant_id and wbes_name:
        query = {"date": yyyy_mm_dd, "$or": [{"wbes_name": wbes_name}, {"wbes_acronym": wbes_name}]}
    elif plant_id and not wbes_name:
        query = {"date": yyyy_mm_dd, "plant_id": plant_id}

    existing = db.db[RAW_DATA_COLLECTION].find_one(query)
    selector = {"_id": existing["_id"]} if existing else {
        "date": yyyy_mm_dd,
        "plant_id": plant_id or "",
        "wbes_name": wbes_name or "",
    }
    base_fields = {
        "date": yyyy_mm_dd,
        "date_dmy": dd_mm_yyyy,
        "plant_id": plant_id or (existing or {}).get("plant_id", ""),
        "wbes_name": wbes_name or (existing or {}).get("wbes_name", ""),
        "wbes_acronym": wbes_name or (existing or {}).get("wbes_acronym", ""),
        "last_updated": datetime.utcnow().isoformat(),
    }
    if set_fields:
        base_fields.update(set_fields)
    db.db[RAW_DATA_COLLECTION].update_one(selector, {"$set": base_fields}, upsert=True)

def get_event_raw_data(db, date_iso: str, plant_id: str = "", wbes_name: str = ""):
    dd_mm_yyyy, yyyy_mm_dd = get_date_formats(date_iso)
    query_parts = []
    if plant_id:
        query_parts.append({"plant_id": plant_id})
    if wbes_name:
        query_parts.extend([{"wbes_name": wbes_name}, {"wbes_acronym": wbes_name}])
    if not query_parts:
        return None
    return db.db[RAW_DATA_COLLECTION].find_one({"date": yyyy_mm_dd, "$or": query_parts})

def get_source_series(raw_doc: Optional[dict], source: str, field: str):
    if not raw_doc:
        return None
    sources = raw_doc.get("sources", {}) or {}
    values = (sources.get(source, {}) or {}).get(field)
    return normalize_series(values) if values is not None else None

def get_source_series_for_timestamp(db, dt_value: datetime, plant_id: str, wbes_name: str, source: str, field: str):
    raw_doc = get_event_raw_data(db, dt_value.date().isoformat(), plant_id=plant_id, wbes_name=wbes_name)
    series = get_source_series(raw_doc, source, field)
    if series is None:
        return None
    return series[block_index_for_time(dt_value, len(series))]

def get_scada_file_series_for_timestamp(db, dt_value: datetime, plant_id: str, wbes_name: str, field: str):
    raw_doc = get_event_raw_data(db, dt_value.date().isoformat(), plant_id=plant_id, wbes_name=wbes_name)
    sources = (raw_doc or {}).get("sources", {}) or {}
    series = (sources.get("scada_file", {}) or {}).get(field)
    if series is None:
        return None
    normalized = normalize_series(series)
    return normalized[block_index_for_time(dt_value, len(normalized))]

def build_database_scada_frame(db, start_dt: datetime, end_dt: datetime, entity_list: list, event_id: Optional[str] = None):
    if event_id:
        event_doc = db.db[EVENT_COLLECTION].find_one({"event_id": event_id}, {"_id": 0})
        event_points = (event_doc or {}).get("data_points") or []
        if event_points:
            point_by_plant = {str(point.get("plant_id")): point for point in event_points}
            first_point = event_points[0]
            timestamps = [
                pd.to_datetime(ts).to_pydatetime()
                for ts in ((first_point.get("series") or {}).get("timestamps") or [])
            ]
            filtered_indexes = [
                idx for idx, ts in enumerate(timestamps)
                if start_dt <= ts <= end_dt
            ]
            scada_dts = [timestamps[idx] for idx in filtered_indexes]
            if scada_dts:
                first_series = first_point.get("series") or {}
                frequency_values = first_series.get("frequency") or []
                df_data_dict = {
                    0: scada_dts,
                    1: [
                        safe_float(frequency_values[idx]) or 0.0
                        for idx in filtered_indexes
                        if idx < len(frequency_values)
                    ],
                }
                if len(df_data_dict[1]) < len(scada_dts):
                    df_data_dict[1].extend([0.0] * (len(scada_dts) - len(df_data_dict[1])))
                headers = ["DATE & TIME", "FREQUENCY"]
                keys = ["DATE & TIME", "FREQUENCY"]
                col_counter = 2
                for entity in entity_list:
                    point = point_by_plant.get(str(entity.get("plant_id")))
                    if not point:
                        continue
                    series = point.get("series") or {}
                    for col_type, db_key in [
                        ("actual", scada_identifier(entity, "actual")),
                        ("schedule", scada_identifier(entity, "schedule")),
                        ("dc", scada_identifier(entity, "dc")),
                    ]:
                        if not db_key:
                            continue
                        values = series.get(col_type) or []
                        df_data_dict[col_counter] = [
                            safe_float(values[idx]) or 0.0
                            for idx in filtered_indexes
                            if idx < len(values)
                        ]
                        if len(df_data_dict[col_counter]) < len(scada_dts):
                            df_data_dict[col_counter].extend([0.0] * (len(scada_dts) - len(df_data_dict[col_counter])))
                        headers.append(str(db_key))
                        keys.append(str(db_key))
                        col_counter += 1
                return pd.DataFrame(df_data_dict), headers, keys, 0, 1, None

    scada_dts = pd.date_range(start=start_dt, end=end_dt, freq='15min').tolist()
    if not scada_dts:
        return None, None, None, None, None, "No timestamps were generated for the selected event range."

    freq_vals = []
    missing_freq_dates = set()
    for ts in scada_dts:
        value = get_source_series_for_timestamp(
            db,
            ts,
            "SYSTEM_FREQUENCY",
            "SYSTEM_FREQUENCY",
            "scada",
            "actual",
        )
        if value is None:
            missing_freq_dates.add(ts.date().isoformat())
            value = 0.0
        freq_vals.append(float(value or 0.0))

    if missing_freq_dates and all(v == 0.0 for v in freq_vals):
        missing = ", ".join(sorted(missing_freq_dates))
        return None, None, None, None, None, f"No frequency data found in database for {missing}."

    df_data_dict = {
        0: scada_dts,
        1: freq_vals,
    }
    headers = ["DATE & TIME", "FREQUENCY"]
    keys = ["DATE & TIME", "FREQUENCY"]
    col_counter = 2

    for e in entity_list:
        pid = e.get("plant_id")
        wbes_name = e.get("wbes_name", "")
        for col_type, db_key in [
            ("actual", scada_identifier(e, "actual")),
            ("schedule", scada_identifier(e, "schedule")),
            ("dc", scada_identifier(e, "dc")),
        ]:
            if not db_key:
                continue
            df_data_dict[col_counter] = [
                get_scada_file_series_for_timestamp(db, ts, pid, wbes_name, col_type) or 0.0
                for ts in scada_dts
            ]
            headers.append(str(db_key))
            keys.append(str(db_key))
            col_counter += 1

    return pd.DataFrame(df_data_dict), headers, keys, 0, 1, None

def save_series_by_date(db, timestamps: list, values: list, plant_id: str, wbes_name: str, field_path: str):
    grouped = {}
    for ts, value in zip(timestamps, values):
        if not isinstance(ts, datetime):
            ts = pd.to_datetime(ts).to_pydatetime()
        date_iso = ts.date().isoformat()
        grouped.setdefault(date_iso, [0.0] * 96)
        grouped[date_iso][block_index_for_time(ts, 96)] = safe_float(value) or 0.0

    for date_iso, series in grouped.items():
        merge_event_raw_data(
            db,
            date_iso,
            plant_id=plant_id,
            wbes_name=wbes_name,
            set_fields={
                field_path: normalize_series(series),
                f"{'.'.join(field_path.split('.')[:2])}.cached_at": datetime.utcnow().isoformat(),
            },
        )

def series_has_data(values):
    return bool(values) and any((safe_float(value) or 0.0) != 0.0 for value in values)

def build_source_status(db, plant_id: str, wbes_name: str, start_dt: datetime, end_dt: datetime):
    dates = get_unique_date_strings(start_dt, end_dt)

    def status(source: str, field: str):
        available_days = 0
        total_points = 0

        for date_iso in dates:
            raw_doc = get_event_raw_data(db, date_iso, plant_id=plant_id, wbes_name=wbes_name)
            sources = (raw_doc or {}).get("sources", {}) or {}
            values = (sources.get(source, {}) or {}).get(field)
            if values is None:
                continue
            normalized = normalize_series(values)
            if series_has_data(normalized):
                available_days += 1
                total_points += sum(1 for value in normalized if (safe_float(value) or 0.0) != 0.0)

        return {
            "available": available_days > 0,
            "days": available_days,
            "total_days": len(dates),
            "points": total_points,
        }

    return {
        "actual": {
            "RTG": status("scada", "actual"),
            "SCADA File": status("scada_file", "actual"),
        },
        "schedule": {
            "RTG": status("rtg", "schedule"),
            "WBES": status("wbes", "schedule"),
            "SCADA File": status("scada_file", "schedule"),
        },
        "dc": {
            "RTG": status("rtg", "dc"),
            "WBES": status("wbes", "dc"),
            "SCADA File": status("scada_file", "dc"),
        },
    }

def build_saved_event_response(db, event_id: str, entity_list: list, start_dt: datetime, end_dt: datetime):
    event_doc = db.db[EVENT_COLLECTION].find_one({"event_id": event_id}, {"_id": 0})
    if not event_doc:
        return None, f"Saved event not found for id {event_id}."

    event_points = event_doc.get("data_points") or []
    if not event_points:
        return None, "Saved event has no merged data points. Please reprocess/upload once and save the event again."

    point_by_plant = {str(point.get("plant_id")): point for point in event_points}
    missing_sources = []
    rows = []
    event_type = normalize_event_type(event_doc.get("event_type"))
    date_strings = get_unique_date_strings(start_dt, end_dt)
    cap_on_bar_by_id = lookup_capacity_on_bar(db, entity_list or event_points, date_strings)

    source_entities = entity_list or event_points

    for entity in source_entities:
        if entity.get("is_frequency"):
            continue
        pid = str(entity.get("plant_id"))
        point = point_by_plant.get(pid)
        if not point:
            missing_sources.append({
                "plant_id": pid,
                "plant_name": entity.get("plant_name") or entity.get("STAGE_NAME") or "",
                "missing": ["actual", "schedule", "dc"],
            })
            continue

        series = point.get("series") or {}
        timestamps = series.get("timestamps") or []
        keep_indexes = []
        for idx, ts in enumerate(timestamps):
            try:
                parsed_ts = pd.to_datetime(ts).to_pydatetime()
                if start_dt <= parsed_ts <= end_dt:
                    keep_indexes.append(idx)
            except Exception:
                pass

        def take(key):
            values = series.get(key) or []
            if not keep_indexes:
                return values
            return [values[idx] for idx in keep_indexes if idx < len(values)]

        filtered_series = {
            "timestamps": take("timestamps"),
            "frequency": take("frequency"),
            "actual": take("actual"),
            "schedule": take("schedule"),
            "dc": take("dc"),
            "deviation": take("deviation"),
        }
        missing = [
            key for key in ["actual", "schedule", "dc"]
            if not series_has_data(filtered_series.get(key))
        ]
        if missing:
            missing_sources.append({
                "plant_id": pid,
                "plant_name": point.get("plant_name") or entity.get("plant_name") or entity.get("STAGE_NAME") or "",
                "missing": missing,
            })

        summary = point.get("summary") or {}
        rtg_pid = entity.get("rtg_plant_id") or pid
        cap_on_bar = summary.get("cap_on_bar") or point.get("cap_on_bar") or cap_on_bar_by_id.get(rtg_pid) or cap_on_bar_by_id.get(pid)
        cap_55 = summary.get("cap_on_bar_55") or ((safe_float(cap_on_bar) or 0) * 0.55 if safe_float(cap_on_bar) is not None else None)
        rows.append({
            "plant_id": pid,
            "plant_name": point.get("plant_name") or entity.get("plant_name") or entity.get("STAGE_NAME") or "",
            "is_state": (point.get("type") == "State") or entity.get("is_state", False),
            "type": "state" if ((point.get("type") == "State") or entity.get("is_state", False)) else "generator",
            "event_type": event_type,
            "capacity": entity.get("stage_installed_capacity") or entity.get("capacity") or 0.0,
            "cap_on_bar": cap_on_bar,
            "cap_on_bar_55": cap_55,
            "avg_capacity_on_bar_pct": summary.get("avg_capacity_on_bar_pct"),
            "actual_source": point.get("actual_source") or entity.get("actual_source") or "RTG",
            "sched_src": point.get("schedule_source") or entity.get("sched_src") or entity.get("schedule_source") or "RTG",
            "dc_src": point.get("dc_source") or entity.get("dc_src") or entity.get("dc_source") or "RTG",
            "wbes_name": entity.get("wbes_name", ""),
            "rtg_plant_id": rtg_pid,
            "scada_key": entity.get("scada_key", ""),
            "scada_header": entity.get("scada_header", ""),
            "scada_schedule_key": entity.get("scada_schedule_key", ""),
            "scada_schedule_header": entity.get("scada_schedule_header", ""),
            "scada_dc_key": entity.get("scada_dc_key", ""),
            "scada_dc_header": entity.get("scada_dc_header", ""),
            "state": entity.get("state_name") or entity.get("state") or "",
            "fuel": entity.get("fuel_type") or entity.get("fuel") or "",
            "owner": entity.get("owner_name") or entity.get("owner") or "",
            "actual": summary.get("actual"),
            "schedule": summary.get("schedule"),
            "dc": summary.get("dc"),
            "deviation": summary.get("deviation"),
            "pct_dc": summary.get("pct_dc"),
            "statistics": summary.get("statistics") or {},
            "series": filtered_series,
            "source_status": build_source_status(db, pid, entity.get("wbes_name", ""), start_dt, end_dt),
            "reason": "Loaded from saved Mongo event.",
        })

    return {
        "success": True,
        "rows": rows,
        "start_time": event_doc.get("start_time"),
        "end_time": event_doc.get("end_time"),
        "event_id": event_id,
        "event_name": event_doc.get("name"),
        "event_type": event_type,
        "from_saved_event": True,
        "missing_sources": missing_sources,
        "logs": [
            f"Loaded historical event from Mongo: {event_doc.get('name')}",
            f"Merged event rows loaded: {len(rows)}",
            f"Missing source groups: {len(missing_sources)}",
        ],
    }, None

def fetch_wbes_schedule_raw(date_str: str, acronyms: list):
    """
    date_str is in DD-MM-YYYY format
    """
    db = MongoService()
    results = []
    missing_acronyms = []
    
    for acr in acronyms:
        try:
            cached = get_event_raw_data(db, date_str, wbes_name=acr)
            cached_schedule = get_source_series(cached, "wbes", "schedule")
            cached_dc = get_source_series(cached, "wbes", "dc")
            if cached_schedule is not None or cached_dc is not None:
                results.append({
                    "Acronym": acr,
                    "NetScheduleSummary": {"TotalNetSchdAmount": cached_schedule or [0.0]*96},
                    "DeclarationList": [{
                        "DeclarationData": {
                            "ThermalDCJsonData": {
                                "SellerInpOnbarAmount": cached_dc or [0.0]*96
                            }
                        }
                    }]
                })
            else:
                legacy_cached = db.db["wbes_schedule_raw"].find_one({"date": date_str, "acronym": acr})
                if legacy_cached:
                    results.append({
                        "Acronym": acr,
                        "NetScheduleSummary": {"TotalNetSchdAmount": normalize_series(legacy_cached.get("schedule"))},
                        "DeclarationList": [{
                            "DeclarationData": {
                                "ThermalDCJsonData": {
                                    "SellerInpOnbarAmount": normalize_series(legacy_cached.get("dc"))
                                }
                            }
                        }]
                    })
                    _, iso_date = get_date_formats(date_str)
                    merge_event_raw_data(
                        db,
                        iso_date,
                        wbes_name=acr,
                        set_fields={
                            "sources.wbes.schedule": normalize_series(legacy_cached.get("schedule")),
                            "sources.wbes.dc": normalize_series(legacy_cached.get("dc")),
                            "sources.wbes.migrated_at": datetime.utcnow().isoformat(),
                        },
                    )
                else:
                    missing_acronyms.append(acr)
        except Exception as cache_err:
            print(f"Error checking cache for WBES {acr}: {cache_err}")
            missing_acronyms.append(acr)
            
    if not missing_acronyms:
        return results

    cfg = db.pipeline_config_collection.find_one({"config_type": "SCHEDULE"})
    if not cfg:
        cfg = {
            "schedule_url": "https://gateway.grid-india.in/POSOCO/reports/1.0/WebAccessAPI/GetUtilityExternalSharedData",
            "schedule_api_key": "6dc32d47-f46a-45c9-afaf-c6ce73c4ca71",
            "schedule_username": "erldc_internal_prod",
            "schedule_password": "ErldcPr0d@052024"
        }
    
    url = f"{cfg['schedule_url']}?apikey={cfg['schedule_api_key']}"
    data = {
        "Date": date_str,
        "SchdRevNo": -1,
        "UserName": cfg["schedule_username"],
        "UtilAcronymList": missing_acronyms,
        "UtilRegionIdList": [1]
    }
    auth = (cfg["schedule_username"], cfg["schedule_password"])
    
    try:
        session = get_legacy_session_no_verify()
        res = session.post(url, json=data, auth=auth, timeout=5)
        status_code = res.status_code
        if status_code == 200:
            resp_body = res.json().get("ResponseBody", {}) or {}
            group_list = resp_body.get("GroupWiseDataList", []) or []
            
            for fsData in group_list:
                acr = fsData.get("Acronym")
                if not acr: continue
                
                totalNetSchdAmount = fsData.get('NetScheduleSummary', {}).get('TotalNetSchdAmount', [0.0]*96)
                
                DCList = [0.0] * 96
                for declaration_entry in fsData.get('DeclarationList', []):
                    decl_data = declaration_entry.get('DeclarationData', {}) or {}
                    for dc_key in ['ThermalDCJsonData', 'GasDCJsonData', 'NuclearDCJsonData', 'HydroDCJsonData']:
                        dc_data = decl_data.get(dc_key)
                        if dc_data and isinstance(dc_data, dict):
                            dc_val = dc_data.get('SellerInpOnbarAmount')
                            if dc_val and isinstance(dc_val, list):
                                DCList = [a + float(b or 0) for a, b in zip(DCList, dc_val)]
                
                _, iso_date = get_date_formats(date_str)
                merge_event_raw_data(
                    db,
                    iso_date,
                    wbes_name=acr,
                    set_fields={
                        "sources.wbes.schedule": normalize_series(totalNetSchdAmount),
                        "sources.wbes.dc": normalize_series(DCList),
                        "sources.wbes.fetched_at": datetime.utcnow().isoformat(),
                    },
                )
                
                results.append(fsData)
            
            log_api_hit("WBES_SCHEDULE", date_str, ",".join(missing_acronyms), url, "success", 200, data_count=len(group_list), data_saved=True)
            return results
        else:
            log_api_hit("WBES_SCHEDULE", date_str, ",".join(missing_acronyms), url, "failed", status_code, f"Response: {res.text[:200]}")
    except Exception as e:
        print("Error fetching WBES schedule:", e)
        log_api_hit("WBES_SCHEDULE", date_str, ",".join(missing_acronyms), url, "failed", 0, str(e))
        
    return results

def fetch_rtg_schedule_raw(date_str: str, plant_id: str):
    """
    date_str is in YYYY-MM-DD format
    """
    db = MongoService()
    
    try:
        cached = get_event_raw_data(db, date_str, plant_id=plant_id)
        cached_schedule = get_source_series(cached, "rtg", "schedule")
        cached_dc = get_source_series(cached, "rtg", "dc")
        if cached_schedule is not None or cached_dc is not None:
            return {
                "schedule": cached_schedule or [0.0]*96,
                "dc": cached_dc or [0.0]*96
            }
        legacy_cached = db.db["rtg_schedule_raw"].find_one({"date": date_str, "plant_id": plant_id})
        if legacy_cached:
            schedule = normalize_series(legacy_cached.get("schedule"))
            dc = normalize_series(legacy_cached.get("dc"))
            merge_event_raw_data(
                db,
                date_str,
                plant_id=plant_id,
                set_fields={
                    "sources.rtg.schedule": schedule,
                    "sources.rtg.dc": dc,
                    "sources.rtg.migrated_at": datetime.utcnow().isoformat(),
                },
            )
            return {"schedule": schedule, "dc": dc}
    except Exception as cache_err:
        print(f"Error checking cache for RTG schedule {plant_id}: {cache_err}")

    rtg_cfg = db.pipeline_config_collection.find_one({"config_type": "RTG"})
    if not rtg_cfg:
        return {}
    
    from services.token_service import TokenService
    try:
        token = TokenService.get_token(
            rtg_cfg["rtg_token_url"],
            rtg_cfg["rtg_username"],
            rtg_cfg["rtg_password"]
        )
    except Exception as e:
        print("Error generating RTG token:", e)
        log_api_hit("RTG_SCHEDULE", date_str, plant_id, rtg_cfg.get("rtg_token_url", ""), "failed", 0, f"Token error: {e}")
        return {}
    
    base_url = get_rtg_schedule_url()
    if not base_url.endswith("/"):
        base_url += "/"
    
    url = f"{base_url}{date_str}/{plant_id}/"
    headers = {
        "Authorization": f"Token {token}",
        "Content-Type": "application/json"
    }
    try:
        session = get_legacy_session_no_verify()
        res = session.get(url, headers=headers, timeout=5)
        status_code = res.status_code
        if status_code == 200:
            res_json = res.json() or {}
            sch = res_json.get("schedule", [0.0]*96)
            dc_val = res_json.get("dc", [0.0]*96)
            
            merge_event_raw_data(
                db,
                date_str,
                plant_id=plant_id,
                set_fields={
                    "sources.rtg.schedule": normalize_series(sch),
                    "sources.rtg.dc": normalize_series(dc_val),
                    "sources.rtg.fetched_at": datetime.utcnow().isoformat(),
                },
            )
            
            log_api_hit("RTG_SCHEDULE", date_str, plant_id, url, "success", 200, data_count=len(sch), data_saved=True)
            return res_json
        else:
            log_api_hit("RTG_SCHEDULE", date_str, plant_id, url, "failed", status_code, f"Response: {res.text[:200]}")
    except Exception as e:
        print(f"Error fetching RTG data for {plant_id} on {date_str}:", e)
        log_api_hit("RTG_SCHEDULE", date_str, plant_id, url, "failed", 0, str(e))
    return {}

def fetch_rtg_scada_raw(date_str: str, plant_id: str):
    """
    date_str is in YYYY-MM-DD format
    """
    db = MongoService()
    
    try:
        cached = get_event_raw_data(db, date_str, plant_id=plant_id)
        cached_actual = get_source_series(cached, "scada", "actual")
        if cached_actual is not None:
            return cached_actual
        legacy_cached = db.db["rtg_scada_raw"].find_one({"date": date_str, "plant_id": plant_id})
        if legacy_cached:
            actual = normalize_series(legacy_cached.get("actual"))
            merge_event_raw_data(
                db,
                date_str,
                plant_id=plant_id,
                set_fields={
                    "sources.scada.actual": actual,
                    "sources.scada.migrated_at": datetime.utcnow().isoformat(),
                },
            )
            return actual
    except Exception as cache_err:
        print(f"Error checking cache for RTG SCADA {plant_id}: {cache_err}")

    rtg_cfg = db.pipeline_config_collection.find_one({"config_type": "RTG"})
    if not rtg_cfg:
        return {}
    
    from services.token_service import TokenService
    try:
        token = TokenService.get_token(
            rtg_cfg["rtg_token_url"],
            rtg_cfg["rtg_username"],
            rtg_cfg["rtg_password"]
        )
    except Exception as e:
        print("Error generating RTG token for SCADA:", e)
        log_api_hit("RTG_SCADA", date_str, plant_id, rtg_cfg.get("rtg_token_url", ""), "failed", 0, f"Token error: {e}")
        return {}
    
    base_url = get_rtg_scada_url()
    if not base_url.endswith("/"):
        base_url += "/"
    
    url = f"{base_url}{date_str}/{plant_id}/"
    headers = {
        "Authorization": f"Token {token}",
        "Content-Type": "application/json"
    }
    try:
        session = get_legacy_session_no_verify()
        res = session.get(url, headers=headers, timeout=5)
        status_code = res.status_code
        if status_code == 200:
            res_json = res.json()
            val_list = []
            if isinstance(res_json, dict):
                for k in ["actual", "data", "scada", "values", "schedule"]:
                    if k in res_json and isinstance(res_json[k], list):
                        val_list = res_json[k]
                        break
            elif isinstance(res_json, list):
                val_list = res_json
            
            merge_event_raw_data(
                db,
                date_str,
                plant_id=plant_id,
                set_fields={
                    "sources.scada.actual": normalize_series(val_list),
                    "sources.scada.fetched_at": datetime.utcnow().isoformat(),
                },
            )
            
            log_api_hit("RTG_SCADA", date_str, plant_id, url, "success", 200, data_count=len(val_list), data_saved=True)
            return res_json
        else:
            log_api_hit("RTG_SCADA", date_str, plant_id, url, "failed", status_code, f"Response: {res.text[:200]}")
    except Exception as e:
        print(f"Error fetching RTG SCADA for {plant_id} on {date_str}:", e)
        log_api_hit("RTG_SCADA", date_str, plant_id, url, "failed", 0, str(e))
    return {}

def get_aligned_schedule_dc(entities: list, start_time: datetime, end_time: datetime):
    # 1. Identify all unique dates in the range
    unique_dates = []
    curr = start_time.date()
    while curr <= end_time.date():
        unique_dates.append(curr)
        curr += timedelta(days=1)
    
    # 2. Pre-fetch WBES data for unique dates
    wbes_acronyms = list(set([
        e.get("wbes_name") for e in entities 
        if e.get("sched_src") == "WBES" or e.get("schedule_source") == "WBES" or 
           e.get("dc_src") == "WBES" or e.get("dc_source") == "WBES" or 
           e.get("is_state")
    ]))
    wbes_acronyms = [a for a in wbes_acronyms if a]
    
    wbes_cache = {}
    from concurrent.futures import ThreadPoolExecutor
    
    def fetch_one_wbes(d):
        dmy = d.strftime('%d-%m-%Y')
        if not wbes_acronyms:
            return [], True
        try:
            group_data = fetch_wbes_schedule_raw(dmy, wbes_acronyms)
            results = []
            if not group_data:
                return [], False
            for fsData in group_data:
                acronym = fsData.get("Acronym")
                totalNetSchdAmount = fsData.get('NetScheduleSummary', {}).get('TotalNetSchdAmount', [0]*96)
                
                NormativeList = [0] * 96
                DCList = [0] * 96
                for declaration_entry in fsData.get('DeclarationList', []):
                    decl_data = declaration_entry.get('DeclarationData', {}) or {}
                    for dc_key in ['ThermalDCJsonData', 'GasDCJsonData', 'NuclearDCJsonData', 'HydroDCJsonData']:
                        dc_data = decl_data.get(dc_key)
                        if dc_data and isinstance(dc_data, dict):
                            norm = dc_data.get('OnbarNormativeAmount')
                            dc_val = dc_data.get('SellerInpOnbarAmount')
                            if norm and isinstance(norm, list):
                                NormativeList = [a + float(b or 0) for a, b in zip(NormativeList, norm)]
                            if dc_val and isinstance(dc_val, list):
                                DCList = [a + float(b or 0) for a, b in zip(DCList, dc_val)]
                results.append(((dmy, acronym), {
                    "schedule": totalNetSchdAmount,
                    "dc": DCList
                }))
            return results, True
        except Exception as e:
            print(f"Error fetching WBES for {dmy}: {e}")
            return [], False
            
    wbes_success = 0
    wbes_failed = 0
    with ThreadPoolExecutor(max_workers=5) as executor:
        wbes_results = list(executor.map(fetch_one_wbes, unique_dates))
        for res_list, ok in wbes_results:
            if ok:
                wbes_success += 1
            else:
                wbes_failed += 1
            for key, val in res_list:
                wbes_cache[key] = val
    
    # 3. Pre-fetch RTG data for unique dates in parallel
    rtg_cache = {}
    
    def fetch_one_rtg(e, d):
        s_src = e.get("sched_src") or e.get("schedule_source") or "RTG"
        d_src = e.get("dc_src") or e.get("dc_source") or "RTG"
        if s_src == "RTG" or d_src == "RTG":
            pid = e.get("rtg_plant_id") or e.get("plant_id")
            if pid:
                iso = d.isoformat()
                try:
                    rtg_data = fetch_rtg_schedule_raw(iso, pid)
                    if rtg_data:
                        return (iso, pid), {
                            "schedule": rtg_data.get("schedule", [0]*96),
                            "dc": rtg_data.get("dc", [0]*96)
                        }, True
                    else:
                        return None, None, False
                except Exception as ex:
                    print(f"Error fetching RTG data for {pid} on {iso}: {ex}")
                    return None, None, False
        return None, None, True

    rtg_tasks = []
    for e in entities:
        for d in unique_dates:
            rtg_tasks.append((e, d))
            
    rtg_success = 0
    rtg_failed = 0
    with ThreadPoolExecutor(max_workers=15) as executor:
        rtg_results = list(executor.map(lambda t: fetch_one_rtg(t[0], t[1]), rtg_tasks))
        for key, val, ok in rtg_results:
            if not ok:
                rtg_failed += 1
            else:
                rtg_success += 1
            if key and val:
                rtg_cache[key] = val
    
    # 4. Generate 1-minute time series for each entity
    dt_index = pd.date_range(start=start_time, end=end_time, freq='1min')
    aligned_data = {}
    
    for e in entities:
        pid = e.get("plant_id")
        sched_src = e.get("sched_src") or e.get("schedule_source") or "RTG"
        dc_src = e.get("dc_src") or e.get("dc_source") or "RTG"
        wbes_name = e.get("wbes_name")
        rtg_pid = e.get("rtg_plant_id") or e.get("plant_id")
        is_frequency = e.get("is_frequency", False)
        
        sched_series = []
        dc_series = []
        
        for t in dt_index:
            if is_frequency:
                sched_series.append(50.0)
                dc_series.append(50.0)
                continue
            
            dmy = t.strftime('%d-%m-%Y')
            iso = t.date().isoformat()
            minute_of_day = t.hour * 60 + t.minute
            
            # --- Get Schedule ---
            sched_val = 0.0
            if sched_src == "WBES" and wbes_name:
                cache_data = wbes_cache.get((dmy, wbes_name), {})
                sch_list = cache_data.get("schedule", [0]*96)
                idx = min(minute_of_day // 15, len(sch_list) - 1)
                sched_val = float(sch_list[idx] or 0)
            elif sched_src == "RTG" and rtg_pid:
                cache_data = rtg_cache.get((iso, rtg_pid), {})
                sch_list = cache_data.get("schedule", [0]*96)
                block_size = 1440 // len(sch_list) if len(sch_list) > 0 else 15
                idx = min(minute_of_day // block_size, len(sch_list) - 1)
                sched_val = float(sch_list[idx] or 0)
            elif sched_src == "Manual":
                sched_val = float(e.get("schedule", 0) or 0)
            
            # --- Get DC ---
            dc_val = 0.0
            if dc_src == "WBES" and wbes_name:
                cache_data = wbes_cache.get((dmy, wbes_name), {})
                dc_list = cache_data.get("dc", [0]*96)
                idx = min(minute_of_day // 15, len(dc_list) - 1)
                dc_val = float(dc_list[idx] or 0)
            elif dc_src == "RTG" and rtg_pid:
                cache_data = rtg_cache.get((iso, rtg_pid), {})
                dc_list = cache_data.get("dc", [0]*96)
                block_size = 1440 // len(dc_list) if len(dc_list) > 0 else 15
                idx = min(minute_of_day // block_size, len(dc_list) - 1)
                dc_val = float(dc_list[idx] or 0)
            elif dc_src == "Manual":
                dc_val = float(e.get("dc", 0) or 0)
            
            sched_series.append(sched_val)
            dc_series.append(dc_val)
            
        aligned_data[pid] = {
            "schedule": sched_series,
            "dc": dc_series
        }
    
    return dt_index, aligned_data, {"success": rtg_success, "failed": rtg_failed}, {"success": wbes_success, "failed": wbes_failed}

SCADA_KEY_FIELDS = {
    "scada_key",
    "scada_schedule_key",
    "scada_dc_key",
}

def normalize_scada_key_value(value):
    if value in [None, ""]:
        return ""
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).strip()
    if text.endswith(".0") and text[:-2].isdigit():
        text = text[:-2]
    if text.isdigit() and not text.startswith("0"):
        return f"0{text}"
    return text

def clean_mapping_value(key, value):
    if value in [None, ""]:
        return ""
    if key in SCADA_KEY_FIELDS:
        return normalize_scada_key_value(value)
    if key in ["plant_id", "STAGE_ID", "rtg_plant_id"]:
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value).strip()
    if isinstance(value, float) and value.is_integer():
        return int(value)
    return value

def parse_scada_file(contents: bytes):
    wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
    sheet_name = "Sheet1" if "Sheet1" in wb.sheetnames else wb.sheetnames[0]
    ws = wb[sheet_name]
    
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 3:
        raise ValueError("SCADA file must have at least 3 rows (header, keys, and values)")
    
    # Dynamically find the row containing date-time header (DATE & TIME) in column 0
    key_row_idx = None
    for idx, row in enumerate(rows):
        if row and row[0] is not None:
            cell_val = str(row[0]).strip().upper()
            if cell_val in ["DATE & TIME", "DATE AND TIME", "DATE &TIME", "DATE/TIME", "DATETIME", "TIME", "DATE_TIME"]:
                key_row_idx = idx
                break
                
    if key_row_idx is None:
        # Fallback to default index if not found
        key_row_idx = 2 if len(rows) > 2 else 1
        
    headers_idx = max(0, key_row_idx - 1)
    
    headers = [str(h).strip() if h is not None else "" for h in rows[headers_idx]]
    keys = [normalize_scada_key_value(k) if k is not None else "" for k in rows[key_row_idx]]
    
    dt_col_idx = 0
    freq_col_idx = 1
    
    data_list = []
    for row in rows[key_row_idx + 1:]:
        if row[dt_col_idx] is None:
            continue
        data_list.append(row)
        
    df_data = pd.DataFrame(data_list)
    
    def parse_dt(val):
        if isinstance(val, datetime):
            return val
        val_str = str(val).strip()
        # Test common formats
        for fmt_str in [
            "%d-%m-%Y %H:%M:%S", "%Y-%m-%d %H:%M:%S", 
            "%d-%m-%Y %H:%M", "%Y-%m-%d %H:%M",
            "%d/%m/%Y %H:%M:%S", "%d/%m/%Y %H:%M",
            "%Y/%m/%d %H:%M:%S", "%Y/%m/%d %H:%M",
            "%d-%b-%Y %H:%M", "%d-%b-%Y %H:%M:%S"
        ]:
            try:
                return datetime.strptime(val_str, fmt_str)
            except ValueError:
                pass
        try:
            return pd.to_datetime(val_str)
        except Exception as e:
            raise ValueError(f"Unknown datetime string format, unable to parse: {val_str}")
        
    df_data[dt_col_idx] = df_data[dt_col_idx].apply(parse_dt)
    df_data = df_data.sort_values(by=dt_col_idx).reset_index(drop=True)
    
    # Identify freq column by key or header
    for idx, (h, k) in enumerate(zip(headers, keys)):
        if k == "04245232" or h.upper() == "FREQUENCY":
            freq_col_idx = idx
            break
            
    return df_data, headers, keys, dt_col_idx, freq_col_idx

def normalize_match_text(value):
    if value in [None, ""]:
        return ""
    text = normalize_scada_key_value(value)
    return " ".join(str(text).strip().upper().split())

def scada_identifier(entity: dict, col_type: str):
    if col_type == "schedule":
        return entity.get("scada_schedule_key") or entity.get("scada_schedule_header") or entity.get("plant_name") or entity.get("STAGE_NAME") or entity.get("plant_id")
    if col_type == "dc":
        return entity.get("scada_dc_key") or entity.get("scada_dc_header") or entity.get("plant_name") or entity.get("STAGE_NAME") or entity.get("plant_id")
    return entity.get("scada_key") or entity.get("scada_header") or entity.get("plant_name") or entity.get("STAGE_NAME") or entity.get("plant_id")

def match_scada_columns(entities: list, headers: list, keys: list):
    matched = {}
    for e in entities:
        pid = e.get("plant_id")
        matched[pid] = {
            "actual": None,
            "schedule": None,
            "dc": None
        }
        for col_type, db_key in [
            ("actual", e.get("scada_key")),
            ("schedule", e.get("scada_schedule_key")),
            ("dc", e.get("scada_dc_key"))
        ]:
            match_candidates = [
                db_key,
                e.get("scada_header") if col_type == "actual" else None,
                e.get("scada_schedule_header") if col_type == "schedule" else None,
                e.get("scada_dc_header") if col_type == "dc" else None,
                e.get("plant_name"),
                e.get("STAGE_NAME"),
                e.get("plant_id"),
                e.get("rtg_plant_id"),
            ]
            candidate_set = {
                normalize_match_text(item)
                for item in match_candidates
                if normalize_match_text(item)
            }
            if not candidate_set:
                continue
            matched_idx = None
            for idx, k in enumerate(keys):
                if normalize_match_text(k) in candidate_set:
                    matched_idx = idx
                    break
            if matched_idx is None:
                for idx, h in enumerate(headers):
                    if normalize_match_text(h) in candidate_set:
                        matched_idx = idx
                        break
            if matched_idx is not None:
                matched[pid][col_type] = matched_idx
    return matched

def resolve_plant_data_series(
    db,
    e,
    scada_dts,
    df_filtered,
    matched_cols,
    keys,
    headers,
    aligned_schedules_dc,
    rtg_scada_cache,
    timestamp_to_idx
):
    pid = e.get("plant_id")
    rtg_pid = e.get("rtg_plant_id") or pid

    sched_src = e.get("sched_src") or e.get("schedule_source") or "RTG"
    dc_src = e.get("dc_src") or e.get("dc_source") or "RTG"
    act_src = e.get("actual_source") or "RTG"

    def column_series(col_idx):
        if col_idx is None:
            return None
        values = []
        for value in df_filtered[col_idx].tolist():
            parsed = safe_float(value)
            values.append(parsed if parsed is not None else 0.0)
        return values

    def has_values(values):
        return bool(values) and any((safe_float(v) or 0.0) != 0.0 for v in values)

    def at_index(values, idx, fallback_idx=None):
        if not values:
            return None
        if fallback_idx is None:
            fallback_idx = idx
        use_idx = fallback_idx if fallback_idx is not None else idx
        use_idx = min(max(int(use_idx), 0), len(values) - 1)
        parsed = safe_float(values[use_idx])
        return parsed if parsed is not None else 0.0

    col_map = matched_cols.get(pid, {})
    scada_act_series = column_series(col_map.get("actual"))
    scada_sch_series = column_series(col_map.get("schedule"))
    scada_dc_series = column_series(col_map.get("dc"))

    aligned = aligned_schedules_dc.get(pid, {}) or {}
    aligned_sch = aligned.get("schedule")
    aligned_dc = aligned.get("dc")

    actual_series = []
    sched_series = []
    dc_series = []
    actual_logs = []
    sched_logs = []
    dc_logs = []

    def add_log(logs, value):
        if value not in logs:
            logs.append(value)

    for idx, t in enumerate(scada_dts):
        yyyy_mm_dd = t.date().isoformat()
        minute_of_day = t.hour * 60 + t.minute
        cache_idx = timestamp_to_idx.get(t)

        actual_val = 0.0
        actual_src = "Missing"
        if act_src == "SCADA":
            if scada_act_series is not None:
                actual_val = at_index(scada_act_series, idx)
                actual_src = "SCADA File"
            else:
                actual_src = "SCADA File (Missing Column)"
        elif act_src == "RTG":
            rtg_scada = (
                rtg_scada_cache.get((yyyy_mm_dd, rtg_pid))
                or rtg_scada_cache.get((yyyy_mm_dd, pid))
            )
            if has_values(rtg_scada):
                block_size = max(1, 1440 // len(rtg_scada))
                actual_val = at_index(rtg_scada, idx, minute_of_day // block_size)
                actual_src = "RTG Portal"
            elif scada_act_series is not None:
                actual_val = at_index(scada_act_series, idx)
                actual_src = "SCADA File (RTG fallback)"
            else:
                actual_src = "Missing (RTG & SCADA unavailable)"
        elif scada_act_series is not None:
            actual_val = at_index(scada_act_series, idx)
            actual_src = "SCADA File (fallback)"

        sched_val = None
        sched_src_log = "Missing"
        if sched_src == "Manual":
            sched_val = safe_float(e.get("schedule")) or 0.0
            sched_src_log = "Manual"
        elif sched_src == "SCADA":
            if scada_sch_series is not None:
                sched_val = at_index(scada_sch_series, idx)
                sched_src_log = "SCADA File"
            else:
                sched_src_log = "SCADA File (Missing Column)"
        elif sched_src in ["WBES", "RTG"]:
            if has_values(aligned_sch):
                sched_val = at_index(aligned_sch, idx, cache_idx)
                sched_src_log = f"{sched_src} Portal"
            elif scada_sch_series is not None:
                sched_val = at_index(scada_sch_series, idx)
                sched_src_log = f"SCADA File ({sched_src} fallback)"
            else:
                sched_src_log = f"Missing ({sched_src} & SCADA unavailable)"

        dc_val = None
        dc_src_log = "Missing"
        if dc_src == "Manual":
            dc_val = safe_float(e.get("dc")) or 0.0
            dc_src_log = "Manual"
        elif dc_src == "SCADA":
            if scada_dc_series is not None:
                dc_val = at_index(scada_dc_series, idx)
                dc_src_log = "SCADA File"
            else:
                dc_src_log = "SCADA File (Missing Column)"
        elif dc_src in ["WBES", "RTG"]:
            if has_values(aligned_dc):
                dc_val = at_index(aligned_dc, idx, cache_idx)
                dc_src_log = f"{dc_src} Portal"
            elif scada_dc_series is not None:
                dc_val = at_index(scada_dc_series, idx)
                dc_src_log = f"SCADA File ({dc_src} fallback)"
            else:
                dc_src_log = f"Missing ({dc_src} & SCADA unavailable)"

        actual_series.append(actual_val)
        sched_series.append(sched_val)
        dc_series.append(dc_val)
        add_log(actual_logs, actual_src)
        add_log(sched_logs, sched_src_log)
        add_log(dc_logs, dc_src_log)

    if all(v is None for v in sched_series):
        sched_series = None
        deviation_series = None
    else:
        deviation_series = [
            actual - (sched if sched is not None else 0.0)
            for actual, sched in zip(actual_series, sched_series)
        ]

    if all(v is None for v in dc_series):
        dc_series = None

    return (
        actual_series,
        sched_series,
        dc_series,
        deviation_series,
        "/".join(actual_logs),
        "/".join(sched_logs),
        "/".join(dc_logs),
    )

# ──────────────────────────────────────────────────────────────
# PLOT GENERATION
# ──────────────────────────────────────────────────────────────

def generate_plot_base64(row_data: dict, start_time: datetime, end_time: datetime):
    import matplotlib.ticker as ticker
    timestamps_str = row_data["series_timestamps"]
    timestamps = [datetime.strptime(ts, '%Y-%m-%d %H:%M:%S') for ts in timestamps_str]
    deviations = row_data.get("series_deviation")
    frequencies = row_data.get("series_frequency", [])
    is_state = row_data["is_state"]
    pname = row_data["plant_name"]
    
    times = np.array(timestamps)
    freqs = np.array(frequencies) if frequencies else np.zeros(len(times))
    
    has_dev = deviations is not None and len(deviations) > 0
    if has_dev:
        devs = np.array(deviations)
        max_abs_dev = max(1.0, float(np.max(np.abs(devs))))
    else:
        devs = np.zeros(len(times))
        max_abs_dev = 10.0
        
    fig, ax1 = plt.subplots(figsize=(12, 5.5), dpi=120)
    
    dev_color = '#10B981' if is_state else '#EF4444'
    freq_color = '#8B5CF6' if is_state else '#3B82F6'
    
    if has_dev:
        ax1.plot(times, devs, color=dev_color, linewidth=1.5, label='Deviation (MW)')
        ax1.set_ylabel('Deviation (MW)', color=dev_color, fontweight='bold', fontsize=9)
        ax1.set_ylim(-max_abs_dev * 1.1, max_abs_dev * 1.1)
    else:
        ax1.set_ylabel('Deviation Not Available (MW)', color='gray', fontweight='bold', fontsize=9)
        ax1.set_ylim(-10, 10)
        
    ax1.set_xlabel('Date-Time', fontweight='bold', fontsize=9)
    ax1.tick_params(axis='y', labelcolor='#475569', labelsize=8)
    ax1.tick_params(axis='x', labelcolor='#475569', labelsize=8)
    
    ax2 = ax1.twinx()
    ax2.plot(times, freqs, color=freq_color, linewidth=1.0, linestyle='--', label='Frequency (Hz)')
    ax2.set_ylabel('Frequency (Hz)', color=freq_color, fontweight='bold', fontsize=9)
    ax2.tick_params(axis='y', labelcolor='#475569', labelsize=8)
    
    ax2.set_ylim(49.4, 50.6)
    
    ax1.axhline(0, color='gray', linestyle='-', linewidth=0.8)
    ax2.axhline(50.0, color='purple', linestyle=':', linewidth=0.8, alpha=0.5)
    
    ax1.grid(True, which='both', linestyle=':', color='lightgray')
    
    ax1.xaxis.set_major_formatter(mdates.DateFormatter('%d-%m-%y %H:%M'))
    ax1.yaxis.set_major_formatter(ticker.FormatStrFormatter('%d'))
    ax2.yaxis.set_major_formatter(ticker.FormatStrFormatter('%.3f'))
    fig.autofmt_xdate()
    
    if has_dev:
        # Low frequency operation shading (< 49.9 Hz)
        freq_low = freqs < 49.9
        if is_state:
            # Over Drawal (Orange): Freq < 49.9 & Dev > 0 (Detoriating)
            orange_fill = freq_low & (devs > 0)
            # Helping Grid (Green): Freq < 49.9 & Dev < 0 (Helping)
            green_fill = freq_low & (devs < 0)
            
            ax1.fill_between(times, 0, devs, where=orange_fill, color='#F97316', alpha=0.35, label='Over Drawal (Orange Shade)')
            ax1.fill_between(times, 0, devs, where=green_fill, color='#10B981', alpha=0.35, label='Under Drawal (Green Shade)')
        else:
            # Under Injection (Orange): Freq < 49.9 & Dev < 0 (Detoriating)
            orange_fill = freq_low & (devs < 0)
            # Helping Grid (Green): Freq < 49.9 & Dev > 0 (Helping)
            green_fill = freq_low & (devs > 0)
            
            ax1.fill_between(times, 0, devs, where=orange_fill, color='#F97316', alpha=0.35, label='Under Injection (Orange Shade)')
            ax1.fill_between(times, 0, devs, where=green_fill, color='#10B981', alpha=0.35, label='Over Injection (Green Shade)')
            
    period_str = f"{start_time.strftime('%d-%m-%y %H:%M')} to {end_time.strftime('%d-%m-%y %H:%M')}"
    plt.title(f"{pname.strip()}: Frequency (Hz) vs Deviation (MW)\nLow Frequency Operation: {period_str}", 
              fontweight='bold', fontsize=11, pad=12)
              
    if not has_dev:
        ann_text = "Schedule data not available;\nDeviation compliance not computed."
    elif is_state:
        ann_text = (
            f"+Ve Dev : Over Drawal (Orange Shade)\n"
            f"% Duration (Freq<49.9 & Dev>0): {safe_float(row_data.get('over_drawal_pct')) or 0.0:.1f}%\n\n"
            f"-Ve Dev : Under Drawal (Green Shade)\n"
            f"% Duration (Freq<49.9 & Dev<0): {safe_float(row_data.get('under_drawal_pct')) or 0.0:.1f}%\n\n"
            f"Max OD = {safe_float(row_data.get('max_od')) or 0.0:.0f} MW\n"
            f"Time: {row_data.get('max_od_time') or '—'}\n"
            f"Freq = {safe_float(row_data.get('max_od_freq') or row_data.get('freq_at_max_od')) or 50.000:.3f} Hz"
        )
    else:
        ann_text = (
            f"+Ve Dev : Over Injection (Green Shade)\n"
            f"% Duration (Freq<49.9 & Dev>0): {safe_float(row_data.get('helping_grid_pct')) or 0.0:.1f}%\n\n"
            f"-Ve Dev : Under Injection (Orange Shade)\n"
            f"% Duration (Freq<49.9 & Dev<0): {safe_float(row_data.get('under_inj_pct')) or 0.0:.1f}%"
        )
        
    props = dict(boxstyle='round', facecolor='white', edgecolor='lightgray', alpha=0.9)
    ax1.text(0.02, 0.98, ann_text, transform=ax1.transAxes, fontsize=8,
             verticalalignment='top', bbox=props, family='monospace')
             
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=120)
    plt.close(fig)
    buf.seek(0)
    img_b64 = base64.b64encode(buf.read()).decode('utf-8')
    return img_b64


@router.get("/export-mapping")
async def export_mapping():
    db = MongoService()
    mappings = list(db.map_collection.find({}, {"_id": 0}))
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Mappings"
    
    headers = [
        "plant_id", "plant_name", "STAGE_ID", "STAGE_NAME", "wbes_name",
        "wbes_acronym", "scada_key", "scada_header", "scada_schedule_key",
        "scada_schedule_header", "scada_dc_key", "scada_dc_header",
        "schedule_source", "dc_source", "actual_source",
        "type", "rtg_plant_id", "is_state", "is_frequency", "stage_installed_capacity"
    ]
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)
        
    for row_num, mapping in enumerate(mappings, 2):
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=row_num, column=col_num)
            val = mapping.get(header, "")
            if val is True:
                cell.value = 1
            elif val is False:
                cell.value = 0
            else:
                cell.value = val
                
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=frequency_mapping.xlsx"}
    )


@router.post("/import-mapping")
async def import_mapping(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents), dtype=str)
        
        # Drop completely empty rows from the Excel
        df = df.dropna(how='all')
        
        errors = []
        required = ["plant_id"]
        for r in required:
            if r not in df.columns:
                errors.append(f"Missing required column: {r}")
                
        if errors:
            return {"success": False, "errors": errors}
            
        if "STAGE_ID" in df.columns:
            dups = df[df.duplicated(["plant_id", "STAGE_ID"], keep=False)]
            duplicate_refs = [
                f"{row.get('plant_id')} / {row.get('STAGE_ID')}"
                for _, row in dups.dropna(subset=["plant_id"]).iterrows()
            ]
            if duplicate_refs:
                errors.append(f"Duplicate plant_id/STAGE_ID references found: {', '.join(map(str, duplicate_refs[:10]))}")
        else:
            dups_id = df[df.duplicated(["plant_id"], keep=False)]["plant_id"].dropna().astype(str).unique().tolist()
            if dups_id:
                errors.append(f"Duplicate plant_id values found: {', '.join(dups_id)}")
            
        for idx, row in df.iterrows():
            r_idx = idx + 2
            if pd.isna(row.get("plant_id")):
                errors.append(f"Row {r_idx}: plant_id is empty")
                
            for s_col in ["schedule_source", "dc_source", "actual_source"]:
                val = row.get(s_col)
                if pd.notna(val) and val not in ["RTG", "WBES", "Manual", "SCADA"]:
                    errors.append(f"Row {r_idx}: Invalid value for {s_col} ('{val}'). Must be one of RTG, WBES, Manual, SCADA")
                    
        if errors:
            return {"success": False, "errors": errors}
            
        db = MongoService()
        new_docs = []
        for _, row in df.iterrows():
            doc = row.to_dict()
            if "plant_name" not in doc or pd.isna(doc.get("plant_name")) or str(doc.get("plant_name")).strip() == "":
                doc["plant_name"] = doc.get("STAGE_NAME") or doc.get("plant_id") or ""
            if "STAGE_NAME" not in doc or pd.isna(doc.get("STAGE_NAME")):
                doc["STAGE_NAME"] = doc.get("plant_name") or ""
            if "STAGE_ID" not in doc or pd.isna(doc.get("STAGE_ID")):
                doc["STAGE_ID"] = doc.get("plant_id") or ""
            cleaned = {}
            for k, v in doc.items():
                if k in ["is_state", "is_frequency"]:
                    cleaned[k] = bool(v) if pd.notna(v) and v not in ["", 0, "0", False] else False
                elif pd.isna(v):
                    cleaned[k] = ""
                elif k == "stage_installed_capacity":
                    cleaned[k] = float(v or 0)
                else:
                    cleaned[k] = clean_mapping_value(k, v)
            new_docs.append(cleaned)
            
        db.db["frequency_mapping"].delete_many({})
        if new_docs:
            db.db["frequency_mapping"].insert_many(new_docs)
        
        return {"success": True, "message": f"Successfully imported {len(new_docs)} mapping records!"}
    except Exception as e:
        return {"success": False, "errors": [str(e)]}


@router.post("/upload-temp-file")
async def upload_temp_file(file: UploadFile = File(...)):
    try:
        os.makedirs("temp_uploads", exist_ok=True)
        file_id = str(uuid.uuid4())
        file_path = os.path.join("temp_uploads", f"{file_id}.xlsx")
        contents = await file.read()
        with open(file_path, "wb") as f:
            f.write(contents)
        return {"success": True, "file_id": file_id}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/process-report-job")
async def create_process_report_job(payload: dict = Body(...)):
    now = datetime.utcnow()
    stale_ids = [
        job_id
        for job_id, job in SSE_REPORT_JOBS.items()
        if (now - job.get("created_at", now)).total_seconds() > 3600
    ]
    for job_id in stale_ids:
        SSE_REPORT_JOBS.pop(job_id, None)

    job_id = str(uuid.uuid4())
    SSE_REPORT_JOBS[job_id] = {
        "created_at": now,
        "payload": {
            "file_id": payload.get("file_id"),
            "start_time": payload.get("start_time"),
            "end_time": payload.get("end_time"),
            "entities": payload.get("entities") or [],
            "event_id": payload.get("event_id") or None,
            "event_type": normalize_event_type(payload.get("event_type")),
        },
    }
    return {"success": True, "job_id": job_id}


@router.get("/process-report-sse")
async def process_report_sse(
    file_id: Optional[str] = None,
    start_time: Optional[str] = None,
    end_time: Optional[str] = None,
    entities: str = "[]",
    job_id: Optional[str] = None,
    event_id: Optional[str] = None
):
    async def sse_generator():
        try:
            if job_id:
                job = SSE_REPORT_JOBS.get(job_id)
                if not job:
                    yield "data: " + json.dumps({"success": False, "error": "Report job expired. Please run again."}) + "\n\n"
                    return
                payload = job.get("payload") or {}
                file_id_from_job = payload.get("file_id")
                start_time_from_job = payload.get("start_time")
                end_time_from_job = payload.get("end_time")
                event_id_from_job = payload.get("event_id")
                event_type_from_job = payload.get("event_type")
                entity_payload = payload.get("entities") or []
                file_id = file_id_from_job
                start_time = start_time_from_job
                end_time = end_time_from_job
                event_id = event_id_from_job
                event_type = normalize_event_type(event_type_from_job)
                entities = json.dumps(entity_payload)
            else:
                event_type = normalize_event_type(None)

            if not file_id or not start_time or not end_time:
                yield "data: " + json.dumps({"success": False, "error": "Missing report job parameters."}) + "\n\n"
                return

            st = datetime.fromisoformat(start_time.replace("Z", ""))
            et = datetime.fromisoformat(end_time.replace("Z", ""))
            entity_list = json.loads(entities)
            cfg = event_config(event_type)

            if file_id == "database" and event_id:
                yield "data: " + json.dumps({"step": 1, "message": "Reading saved historical event from MongoDB..."}) + "\n\n"
            elif file_id == "database":
                yield "data: " + json.dumps({"step": 1, "message": "Reading merged SCADA/frequency data from MongoDB..."}) + "\n\n"
            else:
                yield "data: " + json.dumps({"step": 1, "message": "Reading configuration and uploaded file..."}) + "\n\n"
            
            db = MongoService()

            if file_id == "database":
                if event_id:
                    saved_response, saved_error = build_saved_event_response(db, event_id, entity_list, st, et)
                    if saved_error:
                        yield "data: " + json.dumps({"success": False, "error": saved_error}) + "\n\n"
                        return
                    yield "data: " + json.dumps({"step": 2, "message": "Loaded saved historical event from Mongo."}) + "\n\n"
                    yield "data: " + json.dumps({"complete": True, "result": saved_response}) + "\n\n"
                    return
                df_filtered, headers, keys, dt_col, freq_col, db_error = build_database_scada_frame(db, st, et, entity_list, event_id=event_id)
                if db_error:
                    yield "data: " + json.dumps({"success": False, "error": db_error}) + "\n\n"
                    return
            else:
                file_path = os.path.join("temp_uploads", f"{file_id}.xlsx")
                if not os.path.exists(file_path):
                    yield "data: " + json.dumps({"success": False, "error": "Uploaded file not found"}) + "\n\n"
                    return
                
                with open(file_path, "rb") as f:
                    contents = f.read()
                df_data, headers, keys, dt_col, freq_col = parse_scada_file(contents)
                df_filtered = df_data[(df_data[dt_col] >= st) & (df_data[dt_col] <= et)].reset_index(drop=True)
                
                # Save system frequency actuals to MongoDB
                try:
                    save_series_by_date(
                        db,
                        df_filtered[dt_col].tolist(),
                        df_filtered[freq_col].tolist(),
                        "SYSTEM_FREQUENCY",
                        "SYSTEM_FREQUENCY",
                        "sources.scada.actual",
                    )
                except Exception as freq_save_err:
                    print(f"Error saving system frequency: {freq_save_err}")
            if df_filtered.empty:
                yield "data: " + json.dumps({"success": False, "error": "No SCADA data found in the selected datetime range"}) + "\n\n"
                return

            unique_dates = get_unique_date_strings(st, et)
            cap_on_bar_by_id = lookup_capacity_on_bar(db, entity_list, unique_dates)

            
            yield "data: " + json.dumps({"step": 2, "message": "Checking MongoDB cache for existing actuals and schedules..."}) + "\n\n"
            
            yield "data: " + json.dumps({"step": 3, "message": "Fetching missing data from RTG & WBES in parallel..."}) + "\n\n"
            
            rtg_available = True
            for d in unique_dates:
                doc = db.rtg_dashboard_collection.find_one(
                    {"snapshot_date": d},
                    sort=[("snapshot_time", -1)]
                )
                if doc and doc.get("record_count", 0) > 0:
                    records = doc.get("data", [])
                    has_actuals = any(float(r.get("actual_gen") or r.get("actual_gen_derived") or 0) > 0 for r in records)
                    if not has_actuals:
                        rtg_available = False
                        break
                else:
                    rtg_available = False
                    break

            dt_index, aligned_schedules_dc, rtg_stats, wbes_stats = get_aligned_schedule_dc(entity_list, st, et)

            rtg_scada_success = 0
            rtg_scada_failed = 0
            rtg_scada_cache = {}
            from concurrent.futures import ThreadPoolExecutor
            
            def fetch_one_scada(e, d):
                is_freq = e.get("is_frequency", False)
                if is_freq: return None, None, True
                pid = e.get("rtg_plant_id") or e.get("plant_id")
                if pid:
                    try:
                        scada_data = fetch_rtg_scada_raw(d, pid)
                        val_list = []
                        if isinstance(scada_data, dict):
                            for k in ["actual", "data", "scada", "values", "schedule"]:
                                if k in scada_data and isinstance(scada_data[k], list):
                                    val_list = scada_data[k]
                                    break
                        elif isinstance(scada_data, list):
                            val_list = scada_data
                            
                        parsed_vals = []
                        for item in val_list:
                            if isinstance(item, dict):
                                val = item.get("actual") or item.get("value") or item.get("actual_gen") or item.get("actual_gen_derived") or 0
                                parsed_vals.append(float(val or 0))
                            else:
                                parsed_vals.append(float(item or 0))
                        if parsed_vals:
                            return (d, pid), parsed_vals, True
                        return None, None, False
                    except Exception as exx:
                        print(f"Error fetching RTG SCADA for {pid} on {d}: {exx}")
                        return None, None, False
                return None, None, True

            scada_tasks = []
            for e in entity_list:
                for d in unique_dates:
                    scada_tasks.append((e, d))
                    
            with ThreadPoolExecutor(max_workers=15) as executor:
                scada_results = list(executor.map(lambda t: fetch_one_scada(t[0], t[1]), scada_tasks))
                for key, val, ok in scada_results:
                    if not ok:
                        rtg_scada_failed += 1
                    else:
                        rtg_scada_success += 1
                    if key and val:
                        rtg_scada_cache[key] = val

            yield "data: " + json.dumps({
                "step": 3,
                "message": f"Fetch complete. RTG SCADA (Success: {rtg_scada_success}, Fail: {rtg_scada_failed}), RTG Sched (Success: {rtg_stats['success']}, Fail: {rtg_stats['failed']}), WBES (Success: {wbes_stats['success']}, Fail: {wbes_stats['failed']})"
            }) + "\n\n"

            yield "data: " + json.dumps({"step": 4, "message": "Aligning and resolving multi-source data priorities..."}) + "\n\n"
            
            logs = []
            if rtg_available:
                logs.append(f"RTG SCADA Actuals fetch status: Success={rtg_scada_success} requests | Failed/Timeout={rtg_scada_failed} requests")
                logs.append(f"RTG Schedules fetch status: Success={rtg_stats['success']} requests | Failed/Timeout={rtg_stats['failed']} requests")
                logs.append(f"WBES Schedules fetch status: Success={wbes_stats['success']} requests | Failed/Timeout={wbes_stats['failed']} requests")
            else:
                logs.append("RTG SCADA Actuals: Unavailable (skipped backend fetch)")
                logs.append(f"RTG Schedules fetch status: Success={rtg_stats['success']} requests | Failed/Timeout={rtg_stats['failed']} requests")
                logs.append(f"WBES Schedules fetch status: Success={wbes_stats['success']} requests | Failed/Timeout={wbes_stats['failed']} requests")

            logs.append(f"Report Period: {st.strftime('%d-%m-%Y %H:%M')} to {et.strftime('%d-%m-%Y %H:%M')}")
            logs.append(f"SCADA data range: {df_filtered[dt_col].iloc[0]} to {df_filtered[dt_col].iloc[-1]} ({len(df_filtered)} rows)")
            
            scada_dts = df_filtered[dt_col].tolist()
            scada_freqs = [float(v or 0) for v in df_filtered[freq_col].tolist()]
            
            matched_cols = match_scada_columns(entity_list, headers, keys)
            timestamp_to_idx = {t: idx for idx, t in enumerate(dt_index)}
            processed_rows = []

            for e in entity_list:
                pid = e.get("plant_id")
                pname = e.get("plant_name") or e.get("STAGE_NAME") or ""
                is_state = e.get("is_state", False)
                is_freq = e.get("is_frequency", False)
                if is_freq: continue
                    
                sched_src = e.get("sched_src") or e.get("schedule_source") or "RTG"
                dc_src = e.get("dc_src") or e.get("dc_source") or "RTG"
                act_src = e.get("actual_source") or "RTG"
                plant_type = e.get("type") or ("State" if is_state else "IPP")
                
                wbes_name = e.get("wbes_name", "")
                rtg_pid = e.get("rtg_plant_id") or e.get("plant_id", "")
                
                # Call resolve_plant_data_series helper
                entity_act, entity_sch, entity_dc, entity_dev, actual_log_src, sched_log_src, dc_log_src = resolve_plant_data_series(
                    db, e, scada_dts, df_filtered, matched_cols, keys, headers, aligned_schedules_dc, rtg_scada_cache, timestamp_to_idx
                )
                
                entity_log = f"[{plant_type}] {pname} ({pid}): Actuals -> {actual_log_src} | Schedule -> {sched_log_src} | DC -> {dc_log_src}"
                logs.append(entity_log)

                try:
                    source_fields = {}
                    col_map = matched_cols.get(pid, {})
                    if col_map.get("actual") is not None:
                        save_series_by_date(
                            db,
                            df_filtered[dt_col].tolist(),
                            df_filtered[col_map["actual"]].tolist(),
                            pid,
                            wbes_name,
                            "sources.scada_file.actual",
                        )
                    if col_map.get("schedule") is not None:
                        save_series_by_date(
                            db,
                            df_filtered[dt_col].tolist(),
                            df_filtered[col_map["schedule"]].tolist(),
                            pid,
                            wbes_name,
                            "sources.scada_file.schedule",
                        )
                    if col_map.get("dc") is not None:
                        save_series_by_date(
                            db,
                            df_filtered[dt_col].tolist(),
                            df_filtered[col_map["dc"]].tolist(),
                            pid,
                            wbes_name,
                            "sources.scada_file.dc",
                        )
                except Exception as cache_err:
                    print(f"Error caching SCADA file series for {pid}: {cache_err}")
                
                # Set up fallback write-up in case schedule is missing
                reason_val = e.get("reason", "").strip()
                if entity_sch is None:
                    missing_msg = "Schedule data not available; deviation compliance not computed."
                    if not reason_val:
                        reason_val = missing_msg
                    elif missing_msg not in reason_val:
                        reason_val = f"{reason_val}. {missing_msg}"

                # Save pipeline log to MongoDB
                try:
                    db.pipeline_log_collection.insert_one({
                        "timestamp": datetime.utcnow().isoformat(),
                        "event_type": "plant_resolution",
                        "plant_id": pid,
                        "plant_name": pname,
                        "date": unique_dates[0] if unique_dates else start_time.date().isoformat(),
                        "actual_source": actual_log_src,
                        "schedule_source": sched_log_src,
                        "dc_source": dc_log_src
                    })
                except Exception as db_log_err:
                    print(f"Error logging plant resolution to mongo: {db_log_err}")

                yield "data: " + json.dumps({"step": 5, "message": f"Running deviation and compliance stats for {pname}..."}) + "\n\n"

                has_deviation = entity_dev is not None
                stat_calc = compute_frequency_statistics(scada_freqs, scada_dts, entity_dev, is_state, event_type)
                positive_pct = stat_calc["positive_pct"]
                negative_pct = stat_calc["negative_pct"]
                    
                avg_sched = (sum(entity_sch) / len(entity_sch)) if (entity_sch and len(entity_sch) > 0) else None
                avg_dc = (sum(entity_dc) / len(entity_dc)) if (entity_dc and len(entity_dc) > 0) else None
                avg_act = (sum(entity_act) / len(entity_act)) if (entity_act and len(entity_act) > 0) else 0.0
                avg_dev = (avg_act - avg_sched) if (avg_sched is not None) else None
                cap_on_bar = cap_on_bar_by_id.get(rtg_pid) or cap_on_bar_by_id.get(pid)
                cap_55 = (cap_on_bar * 0.55) if cap_on_bar is not None else None
                avg_capacity_pct = (avg_act / cap_on_bar * 100) if (cap_on_bar and cap_on_bar > 0) else None
                
                row_data = {
                    "plant_id": pid,
                    "plant_name": pname,
                    "is_state": is_state,
                    "event_type": cfg["event_type"],
                    "capacity": e.get("stage_installed_capacity") or e.get("capacity") or 0.0,
                    "cap_on_bar": cap_on_bar,
                    "cap_on_bar_55": cap_55,
                    "avg_capacity_on_bar_pct": avg_capacity_pct,
                    "schedule": avg_sched,
                    "dc": avg_dc,
                    "actual": avg_act,
                    "deviation": avg_dev,
                    "pct_dc": (avg_act / avg_dc * 100) if (avg_dc is not None and avg_dc > 0) else None,
                    "sched_src": sched_src,
                    "dc_src": dc_src,
                    "actual_source": act_src,
                    "type": plant_type,
                    "wbes_name": wbes_name,
                    "rtg_plant_id": rtg_pid,
                    "scada_key": e.get("scada_key"),
                    "scada_header": e.get("scada_header"),
                    "scada_schedule_key": e.get("scada_schedule_key", ""),
                    "scada_schedule_header": e.get("scada_schedule_header", ""),
                    "scada_dc_key": e.get("scada_dc_key", ""),
                    "scada_dc_header": e.get("scada_dc_header", ""),
                    "reason": reason_val,
                    
                    "series_timestamps": [t.strftime('%Y-%m-%d %H:%M:%S') for t in scada_dts],
                    "series_schedule": entity_sch,
                    "series_dc": entity_dc,
                    "series_actual": entity_act,
                    "series_deviation": entity_dev,
                    "series_frequency": scada_freqs,
                    "source_status": build_source_status(db, pid, wbes_name, st, et),
                }
                
                if is_state:
                    if cfg["event_type"] == "high":
                        row_data.update({
                            "max_ud": stat_calc["state_extreme"],
                            "max_ud_time": stat_calc["state_extreme_time"],
                            "max_ud_freq": stat_calc["state_extreme_freq"],
                            "over_drawal_pct": positive_pct,
                            "under_drawal_pct": negative_pct
                        })
                    else:
                        row_data.update({
                            "max_od": stat_calc["state_extreme"],
                            "max_od_time": stat_calc["state_extreme_time"],
                            "max_od_freq": stat_calc["state_extreme_freq"],
                            "over_drawal_pct": positive_pct,
                            "under_drawal_pct": negative_pct
                        })
                else:
                    row_data.update({
                        "under_inj_pct": negative_pct,
                        "helping_grid_pct": positive_pct
                    })
                processed_rows.append(row_data)

            for row in processed_rows:
                is_state = row.get("is_state", False)
                row["series"] = {
                    "timestamps": row.pop("series_timestamps", []),
                    "frequency":  row.pop("series_frequency", []),
                    "deviation":  row.pop("series_deviation", None) or [],
                    "schedule":   row.pop("series_schedule", None) or [],
                    "actual":     row.pop("series_actual", None) or [],
                    "dc":         row.pop("series_dc", None) or [],
                }
                if is_state:
                    row["statistics"] = {
                        "max_od":               row.pop("max_od", None),
                        "max_ud":               row.pop("max_ud", None),
                        "max_od_time":          row.pop("max_od_time", ""),
                        "max_ud_time":          row.pop("max_ud_time", ""),
                        "freq_at_max_od":       row.pop("max_od_freq", None),
                        "freq_at_max_ud":       row.pop("max_ud_freq", None),
                        "od_duration_pct":      row.pop("over_drawal_pct", None),
                        "helping_duration_pct": row.pop("under_drawal_pct", None),
                        "under_inj_pct":        None,
                        "helping_grid_pct":     None,
                    }
                else:
                    row["statistics"] = {
                        "max_od":               None,
                        "max_od_time":          None,
                        "freq_at_max_od":       None,
                        "od_duration_pct":      None,
                        "helping_duration_pct": None,
                        "under_inj_pct":        row.pop("under_inj_pct", 0.0),
                        "helping_grid_pct":     row.pop("helping_grid_pct", 0.0),
                    }
                row["type"] = "state" if is_state else "generator"
                matching_entity = next((item for item in entity_list if item.get("plant_id") == row["plant_id"]), {})
                row["state"] = row.get("state") or (matching_entity.get("state_name") or matching_entity.get("state") or "")
                row["fuel"]  = row.get("fuel")  or (matching_entity.get("fuel_type") or matching_entity.get("fuel") or "")
                row["owner"] = row.get("owner") or (matching_entity.get("owner_name") or matching_entity.get("owner") or "")

            yield "data: " + json.dumps({"step": 6, "message": "Preparing compliance report response..."}) + "\n\n"

            try:
                if file_id != "database" and os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as rm_err:
                print(f"Error removing temp file: {rm_err}")

            response_data = {
                "success": True,
                "event_type": cfg["event_type"],
                "rows": processed_rows,
                "start_time": start_time,
                "end_time": end_time,
                "logs": logs
            }
            yield "data: " + json.dumps({"complete": True, "result": response_data}) + "\n\n"
        except Exception as ex:
            import traceback
            tb_str = traceback.format_exc()
            print(tb_str)
            yield "data: " + json.dumps({"success": False, "error": str(ex), "traceback": tb_str}) + "\n\n"

    return StreamingResponse(sse_generator(), media_type="text/event-stream")

# ──────────────────────────────────────────────────────────────
# REPORT PROCESSING ENDPOINT
# ──────────────────────────────────────────────────────────────

@router.post("/process-report")
async def process_report(
    start_time: str = Form(...),
    end_time: str = Form(...),
    entities: str = Form(...),
    file: Optional[UploadFile] = File(None)
):
    try:
        st = datetime.fromisoformat(start_time.replace("Z", ""))
        et = datetime.fromisoformat(end_time.replace("Z", ""))
        entity_list = json.loads(entities)
        
        db = MongoService()

        if start_time == "database" or (not file and start_time):
            # Fallback handling if query formats are mixed up, let's keep it safe
            pass

        if not file:
            # Running from database
            df_filtered, headers, keys, dt_col, freq_col, db_error = build_database_scada_frame(db, st, et, entity_list)
            if db_error:
                return {"success": False, "error": db_error}
        else:
            contents = await file.read()
            df_data, headers, keys, dt_col, freq_col = parse_scada_file(contents)
            df_filtered = df_data[(df_data[dt_col] >= st) & (df_data[dt_col] <= et)].reset_index(drop=True)
            
            # Save system frequency actuals to MongoDB
            try:
                save_series_by_date(
                    db,
                    df_filtered[dt_col].tolist(),
                    df_filtered[freq_col].tolist(),
                    "SYSTEM_FREQUENCY",
                    "SYSTEM_FREQUENCY",
                    "sources.scada.actual",
                )
            except Exception as freq_save_err:
                print(f"Error saving system frequency: {freq_save_err}")
        unique_dates = get_unique_date_strings(st, et)

        if df_filtered.empty:
            return {"success": False, "error": "No SCADA data found in the selected datetime range"}

        rtg_available = True
        for d in unique_dates:
            doc = db.rtg_dashboard_collection.find_one(
                {"snapshot_date": d},
                sort=[("snapshot_time", -1)]
            )
            if doc and doc.get("record_count", 0) > 0:
                records = doc.get("data", [])
                has_actuals = any(float(r.get("actual_gen") or r.get("actual_gen_derived") or 0) > 0 for r in records)
                if not has_actuals:
                    rtg_available = False
                    break
            else:
                rtg_available = False
                break

        dt_index, aligned_schedules_dc, rtg_stats, wbes_stats = get_aligned_schedule_dc(entity_list, st, et)

        logs = []
        rtg_scada_success = 0
        rtg_scada_failed = 0
        
        rtg_scada_cache = {}
        from concurrent.futures import ThreadPoolExecutor
        
        def fetch_one_scada(e, d):
            is_freq = e.get("is_frequency", False)
            if is_freq: return None, None, True
            pid = e.get("rtg_plant_id") or e.get("plant_id")
            if pid:
                try:
                    scada_data = fetch_rtg_scada_raw(d, pid)
                    val_list = []
                    if isinstance(scada_data, dict):
                        for k in ["actual", "data", "scada", "values", "schedule"]:
                            if k in scada_data and isinstance(scada_data[k], list):
                                val_list = scada_data[k]
                                break
                    elif isinstance(scada_data, list):
                        val_list = scada_data
                        
                    parsed_vals = []
                    for item in val_list:
                        if isinstance(item, dict):
                            val = item.get("actual") or item.get("value") or item.get("actual_gen") or item.get("actual_gen_derived") or 0
                            parsed_vals.append(float(val or 0))
                        else:
                            parsed_vals.append(float(item or 0))
                    if parsed_vals:
                        return (d, pid), parsed_vals, True
                    return None, None, False
                except Exception as exx:
                    print(f"Error fetching RTG SCADA for {pid} on {d}: {exx}")
                    return None, None, False
            return None, None, True

        scada_tasks = []
        for e in entity_list:
            for d in unique_dates:
                scada_tasks.append((e, d))
                
        with ThreadPoolExecutor(max_workers=15) as executor:
            scada_results = list(executor.map(lambda t: fetch_one_scada(t[0], t[1]), scada_tasks))
            for key, val, ok in scada_results:
                if not ok:
                    rtg_scada_failed += 1
                else:
                    rtg_scada_success += 1
                if key and val:
                    rtg_scada_cache[key] = val

        logs.append(f"RTG SCADA Actuals fetch status: Success={rtg_scada_success} requests | Failed/Timeout={rtg_scada_failed} requests")
        logs.append(f"RTG Schedules fetch status: Success={rtg_stats['success']} requests | Failed/Timeout={rtg_stats['failed']} requests")
        logs.append(f"WBES Schedules fetch status: Success={wbes_stats['success']} requests | Failed/Timeout={wbes_stats['failed']} requests")

        logs.append(f"Report Period: {st.strftime('%d-%m-%Y %H:%M')} to {et.strftime('%d-%m-%Y %H:%M')}")
        logs.append(f"SCADA data range: {df_filtered[dt_col].iloc[0]} to {df_filtered[dt_col].iloc[-1]} ({len(df_filtered)} rows)")
        
        scada_dts = df_filtered[dt_col].tolist()
        scada_freqs = [float(v or 0) for v in df_filtered[freq_col].tolist()]
        
        matched_cols = match_scada_columns(entity_list, headers, keys)
        timestamp_to_idx = {t: idx for idx, t in enumerate(dt_index)}
        processed_rows = []

        for e in entity_list:
            pid = e.get("plant_id")
            pname = e.get("plant_name") or e.get("STAGE_NAME") or ""
            is_state = e.get("is_state", False)
            is_freq = e.get("is_frequency", False)
            if is_freq: continue
                
            sched_src = e.get("sched_src") or e.get("schedule_source") or "RTG"
            dc_src = e.get("dc_src") or e.get("dc_source") or "RTG"
            act_src = e.get("actual_source") or "RTG"
            plant_type = e.get("type") or ("State" if is_state else "IPP")
            
            wbes_name = e.get("wbes_name", "")
            rtg_pid = e.get("rtg_plant_id") or e.get("plant_id", "")
            
            # Call resolve_plant_data_series helper
            entity_act, entity_sch, entity_dc, entity_dev, actual_log_src, sched_log_src, dc_log_src = resolve_plant_data_series(
                db, e, scada_dts, df_filtered, matched_cols, keys, headers, aligned_schedules_dc, rtg_scada_cache, timestamp_to_idx
            )
            
            entity_log = f"[{plant_type}] {pname} ({pid}): Actuals -> {actual_log_src} | Schedule -> {sched_log_src} | DC -> {dc_log_src}"
            logs.append(entity_log)

            try:
                col_map = matched_cols.get(pid, {})
                if col_map.get("actual") is not None:
                    save_series_by_date(
                        db,
                        df_filtered[dt_col].tolist(),
                        df_filtered[col_map["actual"]].tolist(),
                        pid,
                        wbes_name,
                        "sources.scada_file.actual",
                    )
                if col_map.get("schedule") is not None:
                    save_series_by_date(
                        db,
                        df_filtered[dt_col].tolist(),
                        df_filtered[col_map["schedule"]].tolist(),
                        pid,
                        wbes_name,
                        "sources.scada_file.schedule",
                    )
                if col_map.get("dc") is not None:
                    save_series_by_date(
                        db,
                        df_filtered[dt_col].tolist(),
                        df_filtered[col_map["dc"]].tolist(),
                        pid,
                        wbes_name,
                        "sources.scada_file.dc",
                    )
            except Exception as cache_err:
                print(f"Error caching SCADA file series for {pid}: {cache_err}")
            
            # Set up fallback write-up in case schedule is missing
            reason_val = e.get("reason", "").strip()
            if entity_sch is None:
                missing_msg = "Schedule data not available; deviation compliance not computed."
                if not reason_val:
                    reason_val = missing_msg
                elif missing_msg not in reason_val:
                    reason_val = f"{reason_val}. {missing_msg}"

            # Save pipeline log to MongoDB
            try:
                db.pipeline_log_collection.insert_one({
                    "timestamp": datetime.utcnow().isoformat(),
                    "event_type": "plant_resolution",
                    "plant_id": pid,
                    "plant_name": pname,
                    "date": unique_dates[0] if unique_dates else start_time.date().isoformat(),
                    "actual_source": actual_log_src,
                    "schedule_source": sched_log_src,
                    "dc_source": dc_log_src
                })
            except Exception as db_log_err:
                print(f"Error logging plant resolution to mongo: {db_log_err}")

            # --- 3. STATISTICS CALCULATIONS ---
            total_count = len(scada_freqs)
            under_inj_count = 0
            helping_grid_count = 0
            over_drawal_count = 0
            under_drawal_count = 0
            
            max_od_val = -999999.0
            max_od_time = None
            max_od_freq = 50.0
            
            has_deviation = entity_dev is not None
            
            if has_deviation:
                for idx, freq in enumerate(scada_freqs):
                    dev_val = entity_dev[idx]
                    if freq < 49.9:
                        if is_state:
                            if dev_val > 0:
                                over_drawal_count += 1
                            else:
                                under_drawal_count += 1
                        else:
                            if dev_val < 0:
                                under_inj_count += 1
                            else:
                                helping_grid_count += 1
                    if is_state:
                        if dev_val > max_od_val:
                            max_od_val = dev_val
                            max_od_time = scada_dts[idx]
                            max_od_freq = freq
                            
            if has_deviation:
                under_inj_pct = (under_inj_count / total_count * 100) if total_count > 0 else 0.0
                helping_grid_pct = (helping_grid_count / total_count * 100) if total_count > 0 else 0.0
                over_drawal_pct = (over_drawal_count / total_count * 100) if total_count > 0 else 0.0
                under_drawal_pct = (under_drawal_count / total_count * 100) if total_count > 0 else 0.0
            else:
                under_inj_pct = None
                helping_grid_pct = None
                over_drawal_pct = None
                under_drawal_pct = None
                
            avg_sched = (sum(entity_sch) / len(entity_sch)) if (entity_sch and len(entity_sch) > 0) else None
            avg_dc = (sum(entity_dc) / len(entity_dc)) if (entity_dc and len(entity_dc) > 0) else None
            avg_act = (sum(entity_act) / len(entity_act)) if (entity_act and len(entity_act) > 0) else 0.0
            avg_dev = (avg_act - avg_sched) if (avg_sched is not None) else None
            
            row_data = {
                "plant_id": pid,
                "plant_name": pname,
                "is_state": is_state,
                "capacity": e.get("stage_installed_capacity") or e.get("capacity") or 0.0,
                "schedule": avg_sched,
                "dc": avg_dc,
                "actual": avg_act,
                "deviation": avg_dev,
                "pct_dc": (avg_act / avg_dc * 100) if (avg_dc is not None and avg_dc > 0) else None,
                "sched_src": sched_src,
                "dc_src": dc_src,
                "actual_source": act_src,
                "type": plant_type,
                "wbes_name": wbes_name,
                "rtg_plant_id": rtg_pid,
                "scada_key": e.get("scada_key"),
                "scada_header": e.get("scada_header"),
                "scada_schedule_key": e.get("scada_schedule_key", ""),
                "scada_schedule_header": e.get("scada_schedule_header", ""),
                "scada_dc_key": e.get("scada_dc_key", ""),
                "scada_dc_header": e.get("scada_dc_header", ""),
                "reason": reason_val,
                
                "series_timestamps": [t.strftime('%Y-%m-%d %H:%M:%S') for t in scada_dts],
                "series_schedule": entity_sch,
                "series_dc": entity_dc,
                "series_actual": entity_act,
                "series_deviation": entity_dev,
                "series_frequency": scada_freqs,
                "source_status": build_source_status(db, pid, wbes_name, st, et),
            }
            
            if is_state:
                row_data.update({
                    "max_od": max_od_val if (has_deviation and max_od_time is not None) else None,
                    "max_od_time": max_od_time.strftime('%d-%m-%y %H:%M') if (has_deviation and max_od_time is not None) else "",
                    "max_od_freq": max_od_freq if (has_deviation and max_od_time is not None) else None,
                    "over_drawal_pct": over_drawal_pct,
                    "under_drawal_pct": under_drawal_pct
                })
            else:
                row_data.update({
                    "under_inj_pct": under_inj_pct,
                    "helping_grid_pct": helping_grid_pct
                })
            processed_rows.append(row_data)
            
        for row in processed_rows:
            is_state = row.get("is_state", False)
            row["series"] = {
                "timestamps": row.pop("series_timestamps", []),
                "frequency":  row.pop("series_frequency", []),
                "deviation":  row.pop("series_deviation", None) or [],
                "schedule":   row.pop("series_schedule", None) or [],
                "actual":     row.pop("series_actual", None) or [],
                "dc":         row.pop("series_dc", None) or [],
            }
            if is_state:
                row["statistics"] = {
                    "max_od":               row.pop("max_od", None),
                    "max_od_time":          row.pop("max_od_time", ""),
                    "freq_at_max_od":       row.pop("max_od_freq", None),
                    "od_duration_pct":      row.pop("over_drawal_pct", None),
                    "helping_duration_pct": row.pop("under_drawal_pct", None),
                    "under_inj_pct":        None,
                    "helping_grid_pct":     None,
                }
            else:
                row["statistics"] = {
                    "max_od":               None,
                    "max_od_time":          None,
                    "freq_at_max_od":       None,
                    "od_duration_pct":      None,
                    "helping_duration_pct": None,
                    "under_inj_pct":        row.pop("under_inj_pct", 0.0),
                    "helping_grid_pct":     row.pop("helping_grid_pct", 0.0),
                }
            # Add entity type
            row["type"] = "state" if is_state else "generator"
            matching_entity = next((item for item in entity_list if item.get("plant_id") == row["plant_id"]), {})
            row["state"] = row.get("state") or (matching_entity.get("state_name") or matching_entity.get("state") or "")
            row["fuel"]  = row.get("fuel")  or (matching_entity.get("fuel_type") or matching_entity.get("fuel") or "")
            row["owner"] = row.get("owner") or (matching_entity.get("owner_name") or matching_entity.get("owner") or "")

        return {
            "success": True,
            "rows": processed_rows,
            "start_time": start_time,
            "end_time": end_time,
            "logs": logs
        }
    except Exception as e:
        import traceback
        tb_str = traceback.format_exc()
        print(tb_str)
        return {
            "success": False,
            "error": str(e),
            "traceback": tb_str
        }



# ──────────────────────────────────────────────────────────────
# EXCEL / PDF / WORD DOWNLOADS
# ──────────────────────────────────────────────────────────────

@router.post("/download-excel")
async def download_excel(payload: dict):
    rows = payload.get("rows", [])
    event_type = normalize_event_type(payload.get("event_type") or (rows[0].get("event_type") if rows else None))
    cfg = event_config(event_type)
    wb = openpyxl.Workbook()

    HDR_FILL   = PatternFill("solid", fgColor="0F172A")
    HDR_FONT   = Font(color="FFFFFF", bold=True, size=10)
    TITLE_FONT = Font(bold=True, size=13, color="0F172A")
    center = Alignment(horizontal="center", vertical="center")
    thin = Border(
        left=Side(style="thin", color="CCCCCC"), right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),  bottom=Side(style="thin", color="CCCCCC"),
    )

    # ── Sheet 1: Executive Summary ──────────────────────────────
    ws_exec = wb.active
    ws_exec.title = "Executive Summary"
    ws_exec.cell(row=1, column=1, value=f"{cfg['label']} Deviation Compliance Report").font = TITLE_FONT
    ws_exec.cell(row=2, column=1, value=payload.get("executive_summary", ""))
    ws_exec.cell(row=4, column=1, value=f"Period: {payload.get('start_time','')} to {payload.get('end_time','')}").font = Font(italic=True)

    # ── Sheet 2: State Summary ──────────────────────────────────
    ws_state = wb.create_sheet(title="State Summary")
    ws_state.cell(row=1, column=1, value="State Drawal Compliance Summary").font = TITLE_FONT
    ws_state.row_dimensions[1].height = 25
    ws_state.row_dimensions[3].height = 20
    state_headers = [
        "State Name", "Schedule Source", "DC (MW)", "Schedule (MW)", "Actual (MW)",
        "Deviation (MW)", "% DC", "Max UD (MW)" if event_type == "high" else "Max OD (MW)",
        "Time of Max UD" if event_type == "high" else "Time of Max OD",
        "Freq at Max UD (Hz)" if event_type == "high" else "Freq at Max OD (Hz)",
        f"% Duration {cfg['threshold_label']} & Dev>0",
        f"% Duration {cfg['threshold_label']} & Dev<0",
        "Reason"
    ]
    for ci, h in enumerate(state_headers, 1):
        c = ws_state.cell(row=3, column=ci, value=h)
        c.fill = HDR_FILL; c.font = HDR_FONT; c.alignment = center; c.border = thin
    state_row_idx = 4
    for r in rows:
      if r.get("is_state"):
        dc_val = get_stat(r, "dc")
        sched_val = get_stat(r, "schedule")
        act_val = get_stat(r, "actual")
        dev_val = get_stat(r, "deviation")
        pct_val = get_stat(r, "pct_dc")
        
        max_od_val = get_stat(r, "max_ud") if event_type == "high" else get_stat(r, "max_od")
        max_time_val = get_stat(r, "max_ud_time") if event_type == "high" else get_stat(r, "max_od_time")
        freq_at_max_val = get_stat(r, "freq_at_max_ud") if event_type == "high" else get_stat(r, "freq_at_max_od")
        od_dur_val = get_stat(r, "od_duration_pct")
        help_dur_val = get_stat(r, "helping_duration_pct")
        
        ws_state.cell(state_row_idx, 1,  r.get("plant_name", "")).border = thin
        ws_state.cell(state_row_idx, 2,  r.get("sched_src", "")).border = thin
        ws_state.cell(state_row_idx, 3,  safe_round(dc_val, 0)).border = thin
        ws_state.cell(state_row_idx, 4,  safe_round(sched_val, 0)).border = thin
        ws_state.cell(state_row_idx, 5,  safe_round(act_val, 0)).border = thin
        ws_state.cell(state_row_idx, 6,  safe_round(dev_val, 0)).border = thin
        ws_state.cell(state_row_idx, 7,  safe_round(pct_val, 0)).border = thin
        ws_state.cell(state_row_idx, 8,  safe_round(max_od_val, 0)).border = thin
        ws_state.cell(state_row_idx, 9,  max_time_val or "-").border = thin
        ws_state.cell(state_row_idx, 10, safe_round(freq_at_max_val, 3)).border = thin
        ws_state.cell(state_row_idx, 11, safe_round(od_dur_val, 1)).border = thin
        ws_state.cell(state_row_idx, 12, safe_round(help_dur_val, 1)).border = thin
        ws_state.cell(state_row_idx, 13, r.get("reason", "")).border = thin
        state_row_idx += 1

    # ── Sheet 3: Generator Summary ──────────────────────────────
    ws_gen = wb.create_sheet(title="Generator Summary")
    ws_gen.cell(row=1, column=1, value="Generator Deviation Compliance Summary").font = TITLE_FONT
    ws_gen.row_dimensions[1].height = 25
    ws_gen.row_dimensions[3].height = 20
    gen_headers = [
        "Plant Name", "State", "Fuel", "Schedule Source", "DC (MW)", "Schedule (MW)",
        "Actual (MW)", "Deviation (MW)", "% DC", "Cap On Bar (MW)", "55% Cap On Bar (MW)", "Actual % Cap On Bar",
        f"% Duration {cfg['threshold_label']} & Dev<0", f"% Duration {cfg['threshold_label']} & Dev>0", "Reason"
    ]
    for ci, h in enumerate(gen_headers, 1):
        c = ws_gen.cell(row=3, column=ci, value=h)
        c.fill = HDR_FILL; c.font = HDR_FONT; c.alignment = center; c.border = thin
    gen_row_idx = 4
    for r in rows:
      if not r.get("is_state"):
        dc_val = get_stat(r, "dc")
        sched_val = get_stat(r, "schedule")
        act_val = get_stat(r, "actual")
        dev_val = get_stat(r, "deviation")
        pct_val = get_stat(r, "pct_dc")
        under_inj_val = get_stat(r, "under_inj_pct")
        helping_grid_val = get_stat(r, "helping_grid_pct")
        
        ws_gen.cell(gen_row_idx, 1,  r.get("plant_name", "")).border = thin
        ws_gen.cell(gen_row_idx, 2,  r.get("state", "")).border = thin
        ws_gen.cell(gen_row_idx, 3,  r.get("fuel", "")).border = thin
        ws_gen.cell(gen_row_idx, 4,  r.get("sched_src", "")).border = thin
        ws_gen.cell(gen_row_idx, 5,  safe_round(dc_val, 0)).border = thin
        ws_gen.cell(gen_row_idx, 6,  safe_round(sched_val, 0)).border = thin
        ws_gen.cell(gen_row_idx, 7,  safe_round(act_val, 0)).border = thin
        ws_gen.cell(gen_row_idx, 8,  safe_round(dev_val, 0)).border = thin
        ws_gen.cell(gen_row_idx, 9,  safe_round(pct_val, 0)).border = thin
        ws_gen.cell(gen_row_idx, 10, safe_round(r.get("cap_on_bar"), 0)).border = thin
        ws_gen.cell(gen_row_idx, 11, safe_round(r.get("cap_on_bar_55"), 0)).border = thin
        ws_gen.cell(gen_row_idx, 12, safe_round(r.get("avg_capacity_on_bar_pct"), 1)).border = thin
        ws_gen.cell(gen_row_idx, 13, safe_round(under_inj_val, 1)).border = thin
        ws_gen.cell(gen_row_idx, 14, safe_round(helping_grid_val, 1)).border = thin
        ws_gen.cell(gen_row_idx, 15, r.get("reason", "")).border = thin
        gen_row_idx += 1

    for ws in [ws_exec, ws_state, ws_gen]:
        for col in ws.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=deviation_compliance_report.xlsx"}
    )



@router.post("/download-pdf")
async def download_pdf(payload: dict):
    try:
        rows = payload.get("rows", [])
        event_type = normalize_event_type(payload.get("event_type") or (rows[0].get("event_type") if rows else None))
        cfg = event_config(event_type)
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=25, leftMargin=25, topMargin=25, bottomMargin=25)
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#022726'),
            spaceAfter=15,
            alignment=1
        )
        section_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#03624C'),
            spaceBefore=10,
            spaceAfter=10
        )
        text_style = ParagraphStyle(
            'NormalText',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=8
        )
        
        story = []
        
        story.append(Paragraph("POWER SYSTEM DEVIATION ANALYSIS REPORT", title_style))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%d-%m-%Y %H:%M')}", text_style))
        story.append(Spacer(1, 10))
        
        for idx, r in enumerate(rows):
            story.append(Paragraph(f"ENTITY: {r.get('plant_name')}", section_style))
            
            plot_b64 = r.get("plot_image")
            if plot_b64:
                try:
                    plot_data = base64.b64decode(plot_b64)
                    plot_buf = io.BytesIO(plot_data)
                    img = Image(plot_buf, width=500, height=225)
                    story.append(img)
                except Exception as img_err:
                    print(f"Skipping invalid PDF plot image for {r.get('plant_name')}: {img_err}")

            if not r.get("is_state") and r.get("capacity_plot_image"):
                try:
                    plot_data = base64.b64decode(r.get("capacity_plot_image"))
                    plot_buf = io.BytesIO(plot_data)
                    img = Image(plot_buf, width=500, height=190)
                    story.append(Spacer(1, 6))
                    story.append(img)
                except Exception as img_err:
                    print(f"Skipping invalid PDF capacity plot image for {r.get('plant_name')}: {img_err}")
                
            if r.get("is_state"):
                max_od_str = safe_format_mw(get_stat(r, 'max_od'))
                max_od_time_str = str(get_stat(r, 'max_od_time') or "—")
                max_od_freq_str = safe_format_hz(get_stat(r, 'max_od_freq'))
                over_drawal_pct_str = safe_format_pct(get_stat(r, 'over_drawal_pct'))
                under_drawal_pct_str = safe_format_pct(get_stat(r, 'under_drawal_pct'))
                
                data = [
                    ["Stat Name", "Value"],
                    ["Max Over Drawal", max_od_str],
                    ["Time of Max OD", max_od_time_str],
                    ["Frequency at Max OD", max_od_freq_str],
                    ["% Duration Over Drawal (Freq<49.9 & Dev>0)", over_drawal_pct_str],
                    ["% Duration Helping Grid (Freq<49.9 & Dev<0)", under_drawal_pct_str]
                ]
            else:
                helping_grid_pct_str = safe_format_pct(get_stat(r, 'helping_grid_pct'))
                under_inj_pct_str = safe_format_pct(get_stat(r, 'under_inj_pct'))
                
                data = [
                    ["Stat Name", "Value"],
                    ["% Duration Helping Grid (Freq<49.9 & Dev>0)", helping_grid_pct_str],
                    ["% Duration Under Injection (Freq<49.9 & Dev<0)", under_inj_pct_str]
                ]
                
            t = Table(data, colWidths=[280, 220])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F172A')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 8),
                ('BOTTOMPADDING', (0,0), (-1,0), 4),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F8FAFC')),
                ('GRID', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
                ('FONTSIZE', (0,1), (-1,-1), 8),
            ]))
            story.append(Spacer(1, 8))
            story.append(t)
            
            if idx < len(rows) - 1:
                story.append(PageBreak())
                
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=deviation_compliance_report.pdf"}
        )
    except Exception as e:
        import traceback
        with open(r"d:\RTG-project\backend\error.log", "w") as f:
            f.write(traceback.format_exc())
        raise e

@router.post("/download-docx")
async def download_docx(payload: dict):
    intro_desc = payload.get("intro_desc", "")
    gen_desc = payload.get("gen_desc", "")
    state_desc = payload.get("state_desc", "")
    rows = payload.get("rows", [])
    
    doc = Document()
    
    title = doc.add_heading("Power System Deviation Analysis Report", level=0)
    title.style.font.color.rgb = docx.shared.RGBColor(2, 39, 38)
    
    doc.add_paragraph(f"Report Date Range: {payload.get('start_time', '')} to {payload.get('end_time', '')}")
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    
    if intro_desc:
        doc.add_heading("1. Executive Summary & General Notes", level=1)
        doc.add_paragraph(intro_desc)
        
    # Generator Section
    doc.add_heading("2. Generator Scheduling Compliance Details", level=1)
    if gen_desc:
        doc.add_paragraph(gen_desc)
        
    gen_rows = [r for r in rows if not r.get("is_state")]
    if gen_rows:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Generator Name'
        hdr_cells[1].text = '% Duration Under Inj (Freq<49.9 & Dev<0)'
        hdr_cells[2].text = '% Duration Helping Grid (Freq<49.9 & Dev>0)'
        
        for r in gen_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(r.get("plant_name"))
            under_inj = get_stat(r, 'under_inj_pct')
            helping_grid = get_stat(r, 'helping_grid_pct')
            row_cells[1].text = safe_format_pct(under_inj)
            row_cells[2].text = safe_format_pct(helping_grid)
            
        doc.add_paragraph() # spacing
        for r in gen_rows:
            doc.add_heading(f"Generator: {r.get('plant_name')}", level=2)
            
            plot_b64 = r.get("plot_image")
            if plot_b64:
                try:
                    plot_data = base64.b64decode(plot_b64)
                    plot_buf = io.BytesIO(plot_data)
                    doc.add_picture(plot_buf, width=docx.shared.Inches(5.8))
                except Exception as img_err:
                    print(f"Skipping invalid DOCX plot image for {r.get('plant_name')}: {img_err}")

            capacity_plot_b64 = r.get("capacity_plot_image")
            if capacity_plot_b64:
                try:
                    plot_data = base64.b64decode(capacity_plot_b64)
                    plot_buf = io.BytesIO(plot_data)
                    doc.add_picture(plot_buf, width=docx.shared.Inches(5.8))
                except Exception as img_err:
                    print(f"Skipping invalid DOCX capacity plot image for {r.get('plant_name')}: {img_err}")
                
            doc.add_paragraph(f"Reason/Comments: {r.get('reason', 'None')}")
            
    # State Section
    doc.add_heading("3. State Drawal Compliance Details", level=1)
    if state_desc:
        doc.add_paragraph(state_desc)
        
    state_rows = [r for r in rows if r.get("is_state")]
    if state_rows:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'State Name'
        hdr_cells[1].text = 'Max OD (MW)'
        hdr_cells[2].text = 'Time of Max OD'
        hdr_cells[3].text = 'Freq at Max OD'
        hdr_cells[4].text = '% Duration Over Drawal'
        hdr_cells[5].text = '% Duration Helping Grid'
        
        for r in state_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(r.get("plant_name"))
            
            max_od = get_stat(r, 'max_od')
            max_od_time = get_stat(r, 'max_od_time')
            max_od_freq = get_stat(r, 'max_od_freq')
            over_drawal = get_stat(r, 'over_drawal_pct')
            under_drawal = get_stat(r, 'under_drawal_pct')
            
            row_cells[1].text = safe_format_mw(max_od)
            row_cells[2].text = str(max_od_time or "—")
            row_cells[3].text = safe_format_hz(max_od_freq)
            row_cells[4].text = safe_format_pct(over_drawal)
            row_cells[5].text = safe_format_pct(under_drawal)
            
        doc.add_paragraph() # spacing
        for r in state_rows:
            doc.add_heading(f"State Drawal: {r.get('plant_name')}", level=2)
            
            plot_b64 = r.get("plot_image")
            if plot_b64:
                try:
                    plot_data = base64.b64decode(plot_b64)
                    plot_buf = io.BytesIO(plot_data)
                    doc.add_picture(plot_buf, width=docx.shared.Inches(5.8))
                except Exception as img_err:
                    print(f"Skipping invalid DOCX plot image for {r.get('plant_name')}: {img_err}")
                
            doc.add_paragraph(f"Reason/Comments: {r.get('reason', 'None')}")
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=deviation_compliance_report.docx"}
    )

# ──────────────────────────────────────────────────────────────
# BACKWARD COMPATIBILITY ENDPOINTS (KEEP EXISTING)
# ──────────────────────────────────────────────────────────────

@router.get("/report-data")
async def get_report_data(date: str = Query(..., description="YYYY-MM-DD")):
    db = MongoService()
    mapping = list(db.map_collection.find({}, {"_id": 0}))
    rtg_snap = db.rtg_dashboard_collection.find_one({}, sort=[("snapshot_time", -1)])
    rtg_by_id = {}
    if rtg_snap and "data" in rtg_snap:
        for p in rtg_snap["data"]:
            pid = p.get("plant_id", "")
            if pid:
                rtg_by_id[pid] = p

    psp_doc = db.psp_collection.find_one({"date": date}, {"_id": 0})
    wbes_schedule_by_acronym = {}
    wbes_dc_by_acronym = {}
    state_load_details = {}
    if psp_doc:
        net_schd = psp_doc.get("Net_schd_list", [])
        for item in net_schd:
            name = item.get("EnergyScheduleTypeName") or item.get("EnergyScheduleSubTypeName") or ""
            amounts = item.get("NetSchdAmount", [])
            total = sum(float(v or 0) for v in amounts)
            if name:
                wbes_schedule_by_acronym[name] = total

        stoa = psp_doc.get("pspSTOADetails1", [])
        for item in stoa:
            acronym = item.get("UNIT_NAME") or item.get("UTILITY_NAME") or ""
            dc_val = float(item.get("DC", 0) or 0)
            if acronym:
                wbes_dc_by_acronym[acronym.upper()] = dc_val

        load_list = psp_doc.get("pspstateloaddetailsER", [])
        for item in load_list:
            st_name = (item.get("STATE_NAME") or "").upper()
            if st_name:
                state_load_details[st_name] = {
                    "schedule": float(item.get("DRAWAL_SCHDULE", 0) or 0),
                    "dc": float(item.get("AVAILABILITY", 0) or 0)
                }
        if state_load_details:
            er_sched = sum(v["schedule"] for v in state_load_details.values())
            er_dc = sum(v["dc"] for v in state_load_details.values())
            state_load_details["ER"] = {"schedule": er_sched, "dc": er_dc}

    rows = []
    for m in mapping:
        plant_name     = m.get("plant_name") or m.get("STAGE_NAME") or ""
        state          = m.get("state_name", "")
        fuel           = m.get("fuel_type", "")
        owner          = m.get("owner_name", "")
        capacity       = float(m.get("stage_installed_capacity") or m.get("installed_capacity") or 0)
        sched_src      = m.get("schedule_source", "RTG")
        dc_src         = m.get("dc_source", "RTG")
        wbes_acronym   = (m.get("wbes_acronym") or "").upper()
        rtg_pid        = m.get("rtg_plant_id") or m.get("plant_id") or ""
        plant_id       = m.get("plant_id", "")
        stage_id       = m.get("STAGE_ID", "")
        scada_key      = m.get("scada_key", "")
        scada_header   = m.get("scada_header", "")
        scada_schedule_key = m.get("scada_schedule_key", "")
        scada_schedule_header = m.get("scada_schedule_header", "")
        scada_dc_key = m.get("scada_dc_key", "")
        scada_dc_header = m.get("scada_dc_header", "")
        is_state_row   = m.get("is_state", False)

        schedule = 0.0
        dc       = 0.0

        if is_state_row:
            if plant_id == "SYSTEM_FREQUENCY":
                schedule = 50.0
                dc       = 50.0
            else:
                if sched_src != "Manual":
                    st_lookup_key = plant_name.upper()
                    if "WEST BENGAL" in st_lookup_key:
                        st_lookup_key = "WEST BENGAL"
                    elif "EASTERN REGION" in st_lookup_key or "ER" in st_lookup_key:
                        st_lookup_key = "ER"
                    details = state_load_details.get(st_lookup_key, {"schedule": 0.0, "dc": 0.0})
                    schedule = details["schedule"]
                    dc = details["dc"]
        else:
            if sched_src == "RTG" or dc_src == "RTG":
                rtg = rtg_by_id.get(rtg_pid) or rtg_by_id.get(plant_id) or {}
                if sched_src == "RTG":
                    schedule = float(rtg.get("schedule", 0) or 0)
                if dc_src == "RTG":
                    dc = float(rtg.get("dc", 0) or 0)
            if sched_src == "WBES" and wbes_acronym:
                schedule = wbes_schedule_by_acronym.get(wbes_acronym, schedule)
            if dc_src == "WBES" and wbes_acronym:
                dc = wbes_dc_by_acronym.get(wbes_acronym, dc)

        rows.append({
            "plant_id":     plant_id,
            "stage_id":     stage_id,
            "plant_name":   plant_name,
            "state":        state,
            "fuel":         fuel,
            "owner":        owner,
            "capacity":     capacity,
            "schedule":     schedule,
            "dc":           dc,
            "actual":       None,
            "deviation":    None,
            "pct_dc":       None,
            "reason":       "",
            "scada_key":    scada_key,
            "scada_header": scada_header,
            "scada_schedule_key": scada_schedule_key,
            "scada_schedule_header": scada_schedule_header,
            "scada_dc_key": scada_dc_key,
            "scada_dc_header": scada_dc_header,
            "sched_src":    sched_src,
            "dc_src":       dc_src,
            "is_state":     is_state_row,
            "is_frequency": m.get("is_frequency", False)
        })

    return {
        "success":  True,
        "date":     date,
        "rows":     rows,
        "wbes_loaded": psp_doc is not None,
        "rtg_loaded":  len(rtg_by_id) > 0,
    }

@router.post("/upload-scada")
async def upload_scada(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        wb = openpyxl.load_workbook(io.BytesIO(contents), read_only=True, data_only=True)
        ws = wb.active

        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 2:
            return {"success": False, "error": "File has no data rows"}

        headers = [str(h).strip() if h is not None else "" for h in rows[0]]

        col_data = {h: [] for h in headers if h}
        for row in rows[1:]:
            for idx, h in enumerate(headers):
                if h and idx < len(row):
                    v = row[idx]
                    if v is not None:
                        try:
                            col_data[h].append(float(v))
                        except (ValueError, TypeError):
                            pass

        result = {}
        for h, vals in col_data.items():
            if vals:
                result[h] = {
                    "latest": vals[-1],
                    "average": round(sum(vals) / len(vals), 2),
                    "total": round(sum(vals), 2),
                    "count": len(vals),
                    "all": vals
                }

        raw_rows = []
        for row in rows[1:]:
            raw_rows.append({
                headers[i]: row[i]
                for i in range(len(headers))
                if headers[i]
            })

        return {
            "success":  True,
            "headers":  [h for h in headers if h],
            "columns":  result,
            "raw_rows": raw_rows[:96],
        }

    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/export-excel")
async def export_excel(payload: dict):
    date  = payload.get("date", "")
    rows  = payload.get("rows", [])
    title = f"Frequency Report — {date}"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Frequency Report"

    HDR_FILL = PatternFill("solid", fgColor="022726")
    HDR_FONT = Font(color="FFFFFF", bold=True, size=10)
    ALT_FILL = PatternFill("solid", fgColor="F0FDF4")
    TITLE_FONT = Font(bold=True, size=13, color="022726")
    center = Alignment(horizontal="center", vertical="center")
    thin = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    ws.merge_cells("A1:K1")
    ws["A1"] = title
    ws["A1"].font = TITLE_FONT
    ws["A1"].alignment = center
    ws.row_dimensions[1].height = 30

    curr_row = 3

    freq_row = next((r for r in rows if r.get("is_frequency")), None)
    if freq_row:
        ws.cell(row=curr_row, column=1, value="SYSTEM FREQUENCY").font = Font(bold=True, size=11, color="022726")
        ws.cell(row=curr_row, column=2, value="Target: 50.00 Hz")
        actual_val = freq_row.get("actual")
        if actual_val is not None:
            ws.cell(row=curr_row, column=3, value=f"Actual: {actual_val:.2f} Hz")
            dev_val = actual_val - 50.0
            ws.cell(row=curr_row, column=4, value=f"Deviation: {dev_val:+.2f} Hz").font = Font(
                bold=True, color="FF0000" if abs(dev_val) >= 0.05 else "166534"
            )
        else:
            ws.cell(row=curr_row, column=3, value="Actual: —")
            ws.cell(row=curr_row, column=4, value="Deviation: —")
        curr_row += 2

    states = [r for r in rows if r.get("is_state") and not r.get("is_frequency")]
    if states:
        ws.cell(row=curr_row, column=1, value="STATE COMPLIANCE DETAILS").font = Font(bold=True, size=11, color="022726")
        curr_row += 1
        
        STATE_HEADERS = [
            "State/Region Name", "Sch.Source", "DC Source", "DC (MU)", "Sched (MU)", "Actual (MU)", "Deviation (MU)", "% DC", "Reason"
        ]
        for ci, h in enumerate(STATE_HEADERS, start=1):
            cell = ws.cell(row=curr_row, column=ci, value=h)
            cell.fill = HDR_FILL
            cell.font = HDR_FONT
            cell.alignment = center
            cell.border = thin
        ws.row_dimensions[curr_row].height = 24
        curr_row += 1
        
        for ri, r in enumerate(states):
            actual = r.get("actual")
            schedule = r.get("schedule") or 0.0
            dc = r.get("dc") or 0.0
            deviation = (actual - schedule) if actual is not None else None
            pct_dc = (actual / dc * 100) if (actual is not None and dc) else None
            
            vals = [
                r.get("plant_name", ""),
                r.get("sched_src", ""),
                r.get("dc_src", ""),
                round(dc, 2),
                round(schedule, 2),
                round(actual, 2) if actual is not None else "",
                round(deviation, 2) if deviation is not None else "",
                f"{pct_dc:.1f}%" if pct_dc is not None else "",
                r.get("reason", "")
            ]
            
            fill = ALT_FILL if ri % 2 == 0 else PatternFill()
            for ci, val in enumerate(vals, start=1):
                cell = ws.cell(row=curr_row, column=ci, value=val)
                cell.alignment = Alignment(vertical="center", horizontal="right" if ci in [4,5,6,7,8] else ("center" if ci in [2,3] else "left"))
                cell.border = thin
                if fill.fgColor.rgb != "00000000":
                    cell.fill = fill
                if ci == 7 and isinstance(val, float):
                    cell.font = Font(color="FF0000" if val < 0 else "166534", bold=True)
            curr_row += 1
        curr_row += 2

    plants = [r for r in rows if not r.get("is_state")]
    if plants:
        ws.cell(row=curr_row, column=1, value="PLANT COMPLIANCE DETAILS").font = Font(bold=True, size=11, color="022726")
        curr_row += 1
        
        PLANT_HEADERS = [
            "Plant Name", "State", "Fuel", "Owner", "Capacity (MW)", "DC (MW)", "Sched (MW)", "Actual (MW)", "Deviation (MW)", "% DC", "Reason"
        ]
        for ci, h in enumerate(PLANT_HEADERS, start=1):
            cell = ws.cell(row=curr_row, column=ci, value=h)
            cell.fill = HDR_FILL
            cell.font = HDR_FONT
            cell.alignment = center
            cell.border = thin
        ws.row_dimensions[curr_row].height = 24
        curr_row += 1
        
        total_capacity = total_dc = total_sched = total_actual = total_dev = 0.0
        
        for ri, r in enumerate(plants):
            actual = r.get("actual")
            schedule = r.get("schedule") or 0.0
            dc = r.get("dc") or 0.0
            deviation = (actual - schedule) if actual is not None else None
            pct_dc = (actual / dc * 100) if (actual is not None and dc) else None
            
            vals = [
                r.get("plant_name", ""),
                r.get("state", ""),
                r.get("fuel", ""),
                r.get("owner", ""),
                r.get("capacity") or 0,
                round(dc, 2),
                round(schedule, 2),
                round(actual, 2) if actual is not None else "",
                round(deviation, 2) if deviation is not None else "",
                f"{pct_dc:.1f}%" if pct_dc is not None else "",
                r.get("reason", "")
            ]
            
            fill = ALT_FILL if ri % 2 == 0 else PatternFill()
            for ci, val in enumerate(vals, start=1):
                cell = ws.cell(row=curr_row, column=ci, value=val)
                cell.alignment = Alignment(vertical="center", horizontal="right" if ci in [5,6,7,8,9,10] else "left")
                cell.border = thin
                if fill.fgColor.rgb != "00000000":
                    cell.fill = fill
                if ci == 9 and isinstance(val, float):
                    cell.font = Font(color="FF0000" if val < 0 else "166534", bold=True)
            
            total_capacity += r.get("capacity") or 0
            total_dc += dc
            total_sched += schedule
            if actual is not None:
                total_actual += actual
                if deviation is not None:
                    total_dev += deviation
            curr_row += 1
            
        ws.cell(row=curr_row, column=1, value="TOTAL").font = Font(bold=True)
        for ci, val in enumerate([
            "", "", "", total_capacity, total_dc,
            total_sched, total_actual, total_dev, "", ""
        ], start=2):
            cell = ws.cell(row=curr_row, column=ci, value=round(val, 2) if isinstance(val, float) else val)
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D1FAE5")
            cell.alignment = Alignment(horizontal="right")
            cell.border = thin

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 11)

    ws.freeze_panes = "B3"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f"frequency_report_{date}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

def get_date_formats(date_str: str):
    if "-" in date_str:
        parts = date_str.split("-")
        if len(parts[0]) == 4: # YYYY-MM-DD
            return f"{parts[2]}-{parts[1]}-{parts[0]}", date_str
        elif len(parts[0]) == 2: # DD-MM-YYYY
            return date_str, f"{parts[2]}-{parts[1]}-{parts[0]}"
    return date_str, date_str

class RawDataSavePayload(BaseModel):
    plant_id: str
    date: str
    wbes_name: Optional[str] = ""
    wbes_schedule: Optional[List[float]] = None
    wbes_dc: Optional[List[float]] = None
    rtg_schedule: Optional[List[float]] = None
    rtg_dc: Optional[List[float]] = None
    actual: Optional[List[float]] = None
    scada_file_actual: Optional[List[float]] = None
    scada_file_schedule: Optional[List[float]] = None
    scada_file_dc: Optional[List[float]] = None
    
    # Old fields for backward compatibility
    source: Optional[str] = ""
    schedule: Optional[List[float]] = None
    dc: Optional[List[float]] = None

class FrequencyEventPayload(BaseModel):
    name: str
    start_time: str
    end_time: str
    event_type: Optional[str] = "low"
    notes: Optional[str] = ""
    details: Optional[List[dict]] = None
    data_points: Optional[List[dict]] = None

@router.get("/raw-data")
def get_raw_data(plant_id: str, date: str, source: Optional[str] = "", wbes_name: Optional[str] = ""):
    db = MongoService()
    try:
        dd_mm_yyyy, yyyy_mm_dd = get_date_formats(date)
        acr = wbes_name or plant_id

        raw_doc = get_event_raw_data(db, yyyy_mm_dd, plant_id=plant_id, wbes_name=acr)

        legacy_wbes = db.db["wbes_schedule_raw"].find_one({"date": dd_mm_yyyy, "acronym": acr})
        legacy_rtg = db.db["rtg_schedule_raw"].find_one({"date": yyyy_mm_dd, "plant_id": plant_id})
        legacy_scada = db.db["rtg_scada_raw"].find_one({"date": yyyy_mm_dd, "plant_id": plant_id})

        wbes_schedule = get_source_series(raw_doc, "wbes", "schedule")
        wbes_dc = get_source_series(raw_doc, "wbes", "dc")
        rtg_schedule = get_source_series(raw_doc, "rtg", "schedule")
        rtg_dc = get_source_series(raw_doc, "rtg", "dc")
        actual = get_source_series(raw_doc, "scada", "actual")
        scada_file_actual = get_source_series(raw_doc, "scada_file", "actual")
        scada_file_schedule = get_source_series(raw_doc, "scada_file", "schedule")
        scada_file_dc = get_source_series(raw_doc, "scada_file", "dc")

        wbes_schedule = wbes_schedule if wbes_schedule is not None else normalize_series((legacy_wbes or {}).get("schedule"))
        wbes_dc = wbes_dc if wbes_dc is not None else normalize_series((legacy_wbes or {}).get("dc"))
        rtg_schedule = rtg_schedule if rtg_schedule is not None else normalize_series((legacy_rtg or {}).get("schedule"))
        rtg_dc = rtg_dc if rtg_dc is not None else normalize_series((legacy_rtg or {}).get("dc"))
        actual = actual if actual is not None else normalize_series((legacy_scada or {}).get("actual"))
        
        # Backwards compatible schedule and dc
        schedule = wbes_schedule if source == "WBES" else rtg_schedule
        dc = wbes_dc if source == "WBES" else rtg_dc
        
        return {
            "success": True,
            "wbes_schedule": wbes_schedule,
            "wbes_dc": wbes_dc,
            "rtg_schedule": rtg_schedule,
            "rtg_dc": rtg_dc,
            "actual": actual,
            "scada_file_actual": scada_file_actual if scada_file_actual is not None else [0.0] * 96,
            "scada_file_schedule": scada_file_schedule if scada_file_schedule is not None else [0.0] * 96,
            "scada_file_dc": scada_file_dc if scada_file_dc is not None else [0.0] * 96,
            "schedule": schedule,
            "dc": dc
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/raw-data")
def save_raw_data(payload: RawDataSavePayload):
    db = MongoService()
    try:
        dd_mm_yyyy, yyyy_mm_dd = get_date_formats(payload.date)
        acr = payload.wbes_name or payload.plant_id
        
        # Determine values to save (supporting both new lists and old list fallbacks)
        wbes_schedule = payload.wbes_schedule if payload.wbes_schedule is not None else (payload.schedule if payload.source == "WBES" else None)
        wbes_dc = payload.wbes_dc if payload.wbes_dc is not None else (payload.dc if payload.source == "WBES" else None)
        
        rtg_schedule = payload.rtg_schedule if payload.rtg_schedule is not None else (payload.schedule if payload.source != "WBES" else None)
        rtg_dc = payload.rtg_dc if payload.rtg_dc is not None else (payload.dc if payload.source != "WBES" else None)
        
        set_fields = {}
        if wbes_schedule is not None or wbes_dc is not None:
            if wbes_schedule is not None:
                set_fields["sources.wbes.schedule"] = normalize_series(wbes_schedule)
            if wbes_dc is not None:
                set_fields["sources.wbes.dc"] = normalize_series(wbes_dc)
            set_fields["sources.wbes.edited_at"] = datetime.utcnow().isoformat()
            
        if rtg_schedule is not None or rtg_dc is not None:
            if rtg_schedule is not None:
                set_fields["sources.rtg.schedule"] = normalize_series(rtg_schedule)
            if rtg_dc is not None:
                set_fields["sources.rtg.dc"] = normalize_series(rtg_dc)
            set_fields["sources.rtg.edited_at"] = datetime.utcnow().isoformat()
            
        if payload.actual is not None:
            set_fields["sources.scada.actual"] = normalize_series(payload.actual)
            set_fields["sources.scada.edited_at"] = datetime.utcnow().isoformat()

        if (
            payload.scada_file_actual is not None
            or payload.scada_file_schedule is not None
            or payload.scada_file_dc is not None
        ):
            if payload.scada_file_actual is not None:
                set_fields["sources.scada_file.actual"] = normalize_series(payload.scada_file_actual)
            if payload.scada_file_schedule is not None:
                set_fields["sources.scada_file.schedule"] = normalize_series(payload.scada_file_schedule)
            if payload.scada_file_dc is not None:
                set_fields["sources.scada_file.dc"] = normalize_series(payload.scada_file_dc)
            set_fields["sources.scada_file.edited_at"] = datetime.utcnow().isoformat()

        merge_event_raw_data(
            db,
            yyyy_mm_dd,
            plant_id=payload.plant_id,
            wbes_name=acr,
            set_fields=set_fields,
        )
            
        return {
            "success": True,
            "message": f"Raw data updated in {RAW_DATA_COLLECTION} successfully",
            "collection": RAW_DATA_COLLECTION
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.get("/available-dates")
def get_available_dates():
    try:
        db = MongoService()
        dates = db.db[RAW_DATA_COLLECTION].distinct("date")
        # filter out empty or null dates and sort descending
        valid_dates = sorted([d for d in dates if d], reverse=True)
        events = list(db.db[EVENT_COLLECTION].find({}, {"_id": 0, "data_points": 0}).sort("start_time", -1))
        return {"success": True, "dates": valid_dates, "events": events}
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.get("/events")
def list_frequency_events():
    try:
        db = MongoService()
        events = list(db.db[EVENT_COLLECTION].find({}, {"_id": 0, "data_points": 0}).sort("start_time", -1))
        return {"success": True, "events": events}
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/events")
def save_frequency_event(payload: FrequencyEventPayload):
    try:
        start_dt = datetime.fromisoformat(payload.start_time.replace("Z", ""))
        end_dt = datetime.fromisoformat(payload.end_time.replace("Z", ""))
        if end_dt < start_dt:
            return {"success": False, "error": "Event end time cannot be before start time."}

        name = payload.name.strip()
        if not name:
            name = f"Frequency Event {start_dt.strftime('%d-%m-%Y %H:%M')}"

        db = MongoService()
        date_span = get_unique_date_strings(start_dt, end_dt)
        event_id = str(uuid.uuid4())
        event_type = normalize_event_type(payload.event_type)
        doc = {
            "event_id": event_id,
            "name": name,
            "event_type": event_type,
            "start_time": start_dt.isoformat(timespec="minutes"),
            "end_time": end_dt.isoformat(timespec="minutes"),
            "dates": date_span,
            "duration_minutes": int((end_dt - start_dt).total_seconds() // 60) + 1,
            "notes": payload.notes or "",
            "details": payload.details or [],
            "data_points": payload.data_points or [],
            "single_event_record": True,
            "updated_at": datetime.utcnow().isoformat(),
        }
        db.db[EVENT_COLLECTION].update_one(
            {
                "name": name,
                "start_time": doc["start_time"],
                "end_time": doc["end_time"],
            },
            {
                "$set": doc,
                "$setOnInsert": {"created_at": datetime.utcnow().isoformat()},
            },
            upsert=True,
        )
        saved = db.db[EVENT_COLLECTION].find_one(
            {"name": name, "start_time": doc["start_time"], "end_time": doc["end_time"]},
            {"_id": 0},
        )
        return {"success": True, "event": saved}
    except Exception as e:
        return {"success": False, "error": str(e)}

class ResyncSourcePayload(BaseModel):
    start_time: str
    end_time: str
    source: str # "SCADA" | "WBES" | "RTG"
    entities: List[dict]

@router.post("/resync-source")
def resync_source(payload: ResyncSourcePayload):
    try:
        st = datetime.fromisoformat(payload.start_time.replace("Z", ""))
        et = datetime.fromisoformat(payload.end_time.replace("Z", ""))
        
        unique_dates = []
        curr = st.date()
        while curr <= et.date():
            unique_dates.append(curr)
            curr += timedelta(days=1)
            
        db = MongoService()
        import concurrent.futures
        
        resync_count = 0
        
        if payload.source == "WBES":
            wbes_acronyms = list(set([
                e.get("wbes_name") for e in payload.entities
                if e.get("wbes_name")
            ]))
            for d in unique_dates:
                dmy = d.strftime('%d-%m-%Y')
                # Force fetch and overwrite cache
                fetch_wbes_schedule_raw(dmy, wbes_acronyms)
                resync_count += len(wbes_acronyms)
                
        elif payload.source == "RTG":
            rtg_pids = list(set([
                e.get("rtg_plant_id") or e.get("plant_id")
                for e in payload.entities
                if e.get("rtg_plant_id") or e.get("plant_id")
            ]))
            for d in unique_dates:
                iso = d.isoformat()
                for pid in rtg_pids:
                    # Force fetch and overwrite cache
                    fetch_rtg_schedule_raw(iso, pid)
                    resync_count += 1
                    
        elif payload.source == "SCADA":
            pids = list(set([
                e.get("rtg_plant_id") or e.get("plant_id")
                for e in payload.entities
                if e.get("rtg_plant_id") or e.get("plant_id")
            ]))
            for d in unique_dates:
                iso = d.isoformat()
                for pid in pids:
                    # Force fetch and overwrite cache
                    fetch_rtg_scada_raw(iso, pid)
                    resync_count += 1
                    
        return {
            "success": True,
            "message": f"Successfully resynced {resync_count} {payload.source} records.",
            "resync_count": resync_count
        }
    except Exception as e:
        return {"success": False, "error": str(e)}
